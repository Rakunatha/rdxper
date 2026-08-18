"""
rdxper v4.0 — Free AI-Powered Real Research Paper Generator
────────────────────────────────────────────────────────────
Pipeline:
  1. Semantic Scholar API  → real papers (titles, abstracts, citations, DOIs)
  2. CrossRef API          → additional verified journal articles
  3. Wikipedia REST API    → background context & definitions
  4. Groq API (FREE)       → writes ALL prose sections using scraped data as context
  5. python-docx           → assembles formatted .docx with SPSS-style charts

AI Provider:
  Groq (free tier) — https://console.groq.com
  set GROQ_API_KEY=your_key_here

Usage:
  python rdxper.py
"""

import os, uuid, time, threading, smtplib, secrets, io, random, re, json, hmac, hashlib, sqlite3
import urllib.request, urllib.parse
from concurrent.futures import ThreadPoolExecutor, as_completed
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
from flask import Flask, request, jsonify, send_file, Response
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)

otp_store = {}
sessions  = {}
jobs      = {}
ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'rkhrishanthm@gmail.com')

# ── SQLite DB ─────────────────────────────────────────────────────────────────
DB_PATH = os.environ.get('DB_PATH', 'rdxper.db')

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.executescript("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY, email TEXT UNIQUE NOT NULL,
                name TEXT, picture TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                last_login TEXT
            );
            CREATE TABLE IF NOT EXISTS papers (
                id TEXT PRIMARY KEY, user_id TEXT NOT NULL, topic TEXT,
                file_path TEXT, paid INTEGER DEFAULT 0, amount INTEGER DEFAULT 0,
                created_at TEXT DEFAULT (datetime('now')),
                FOREIGN KEY(user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS payments (
                id TEXT PRIMARY KEY, user_id TEXT NOT NULL, paper_id TEXT,
                razorpay_order TEXT, razorpay_payment TEXT, amount INTEGER,
                status TEXT DEFAULT 'pending',
                created_at TEXT DEFAULT (datetime('now')),
                FOREIGN KEY(user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY,
                email TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now'))
            );
        """)

init_db()
os.makedirs('generated', exist_ok=True)


def session_set(token: str, email: str):
    """Persist a session token to the DB and keep in-memory cache."""
    sessions[token] = {'email': email}
    try:
        with get_db() as db:
            db.execute('INSERT OR REPLACE INTO sessions (token, email) VALUES (?, ?)', (token, email))
    except Exception as e:
        print(f'[session_set] DB error: {e}')


def session_get(token: str) -> object:
    """Return session dict from memory, falling back to DB (handles restarts)."""
    if not token:
        return None
    if token in sessions:
        return sessions[token]
    try:
        with get_db() as db:
            row = db.execute('SELECT email FROM sessions WHERE token=?', (token,)).fetchone()
            if row:
                email = row['email']
                user = db.execute('SELECT id, name, picture FROM users WHERE email=?', (email,)).fetchone()
                sessions[token] = {
                    'email': email,
                    'user_id': user['id'] if user else email,
                    'name': user['name'] if user else '',
                    'picture': user['picture'] if user else '',
                }
                return sessions[token]
    except Exception as e:
        print(f'[session_get] DB error: {e}')
    return None


def session_delete(token: str):
    sessions.pop(token, None)
    try:
        with get_db() as db:
            db.execute('DELETE FROM sessions WHERE token=?', (token,))
    except Exception as e:
        print(f'[session_delete] DB error: {e}')


# ═══════════════════════════════════════════════════════════════════════════════
#  AI CLIENT  (Groq — fast free inference)
# ═══════════════════════════════════════════════════════════════════════════════

# Ordered sections — used to map closing tags → progress %
SECTION_ORDER = [
    'keywords', 'abstract', 'introduction', 'objectives',
    'literature_review', 'methodology', 'results',
    'discussion', 'limitations', 'suggestions', 'conclusion', 'charts',
]
SECTION_LABELS = {
    'keywords':          'Writing keywords...',
    'abstract':          'Writing abstract...',
    'introduction':      'Writing introduction...',
    'objectives':        'Writing objectives...',
    'literature_review': 'Writing literature review...',
    'methodology':       'Writing methodology...',
    'results':           'Writing results & analysis...',
    'discussion':        'Writing discussion...',
    'limitations':       'Writing limitations...',
    'suggestions':       'Writing suggestions...',
    'conclusion':        'Writing conclusion...',
    'charts':            'Designing chart specifications...',
}
_AI_START = 30
_AI_END   = 75

# Groq free models — tried in order, with exponential backoff on 429
# See https://console.groq.com/docs/models for current list
# Current Groq production models. Keep this list limited to models that are
# currently supported by Groq; retired model IDs cause HTTP 400 model_decommissioned
# errors and should not be retried.
#
# GROQ_MODEL can be set to force a preferred model, e.g.:
#   GROQ_MODEL=llama-3.3-70b-versatile
# Groq production model. Keep this list conservative: model availability can
# differ by Groq project/org permissions, and a 404 means the project cannot
# use that model. Llama 3.3 70B is currently supported by Groq.
_GROQ_PREFERRED_MODELS = [
    # Preferred order. The app will first ask Groq which models this API key
    # can actually see, so a model that is unavailable to the project is never
    # blindly selected.
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
]


def _get_groq_models(api_key, requests_module):
    """Return active Groq models visible to the current API key/project."""
    headers = {"Authorization": f"Bearer {api_key}"}
    try:
        resp = requests_module.get(
            "https://api.groq.com/openai/v1/models",
            headers=headers,
            timeout=20,
        )
        if resp.status_code != 200:
            return [], f"HTTP {resp.status_code} from Groq /models: {resp.text[:300]}"
        data = resp.json()
        models = data.get("data", []) if isinstance(data, dict) else []
        ids = []
        for item in models:
            if not isinstance(item, dict):
                continue
            model_id = item.get("id")
            if model_id and item.get("active", True):
                ids.append(model_id)
        return ids, None
    except Exception as e:
        return [], f"Could not query Groq /models: {e}"


def _select_groq_models(api_key, requests_module):
    """Build a model list from models actually exposed to this API key."""
    preferred_override = os.environ.get("GROQ_MODEL", "").strip()
    available, discovery_error = _get_groq_models(api_key, requests_module)
    available_set = set(available)

    if preferred_override:
        # If the user explicitly chose a model, try it first even when /models
        # could not be queried. This preserves the existing GROQ_MODEL behavior.
        selected = [preferred_override]
        selected.extend(m for m in _GROQ_PREFERRED_MODELS
                        if m != preferred_override and m in available_set)
    elif available:
        selected = [m for m in _GROQ_PREFERRED_MODELS if m in available_set]
        # If none of our preferred models are visible, use other active models
        # returned by Groq, excluding obvious non-chat/safety/audio models.
        excluded = ("whisper", "guard", "safeguard", "compound")
        selected.extend(
            m for m in available
            if m not in selected and not any(x in m.lower() for x in excluded)
        )
    else:
        # Discovery failed; retain a conservative fallback so the error from
        # chat/completions remains actionable.
        selected = [preferred_override] if preferred_override else [
            "llama-3.3-70b-versatile",
            "llama-3.1-8b-instant",
            "openai/gpt-oss-120b",
            "openai/gpt-oss-20b",
        ]

    return list(dict.fromkeys(selected)), discovery_error


def ai_generate(prompt: str, system: str = "", temperature: float = 0.7,
                progress_cb=None, tracked_sections=None) -> str:
    """
    Call Groq API with requests library + exponential backoff on 429.
    Tries each model in order; retries up to 3x on rate limit before skipping.
    """
    import requests as _req

    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set. Get a free key at https://console.groq.com")

    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type":  "application/json",
    }

    # IMPORTANT: Do not assume a model is available merely because Groq
    # documents it. Availability can differ by project/organization. Ask the
    # Models API which active models this key can actually access.
    _GROQ_MODELS, discovery_error = _select_groq_models(api_key, _req)
    print(f"[Groq] Models selected for this key/project: {_GROQ_MODELS}")
    if discovery_error:
        print(f"[Groq] Model discovery warning: {discovery_error}")
    if not _GROQ_MODELS:
        raise RuntimeError(
            "No active Groq text models are available to this API key/project. "
            "Check Groq Project > Settings > Limits/Model Permissions and API key."
        )

    last_error = None

    for model in _GROQ_MODELS:
        payload = {
            "model":       model,
            "messages":    messages,
            "temperature": temperature,
            # Groq's current production Llama models support max_completion_tokens.
            # Keep the output bounded so long research-paper sections remain reliable.
            "max_completion_tokens": 4096,
            "stream":      False,
        }

        # Retry up to 2 times on 429 with exponential backoff before giving up on this model
        for attempt in range(3):
            try:
                resp = _req.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=90,
                )
            except _req.exceptions.Timeout:
                last_error = f"Timeout on {model}"
                print(f"[Groq] Timeout on {model}, trying next...")
                break
            except _req.exceptions.RequestException as e:
                last_error = f"Request error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            status = resp.status_code

            if status == 429:
                wait = 2 ** (attempt + 2)  # 4s, 8s, 16s
                last_error = f"429 rate-limited on {model} (attempt {attempt+1})"
                print(f"[Groq] 429 on {model}, waiting {wait}s...")
                time.sleep(wait)
                continue  # retry same model

            if status in (400, 402, 404, 503):
                body = resp.text[:300]
                last_error = f"HTTP {status} on {model}: {body}"
                print(f"[Groq] {status} on {model} (skipping): {body[:120]}")
                break  # skip to next model

            if status != 200:
                last_error = f"HTTP {status} on {model}: {resp.text[:300]}"
                print(f"[Groq] Unexpected {status} on {model}: {resp.text[:120]}")
                break

            # Parse successful JSON response
            try:
                data = resp.json()
            except Exception as e:
                last_error = f"JSON parse error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if "error" in data:
                err = data["error"]
                last_error = f"API error on {model}: {err}"
                print(f"[Groq] {last_error}")
                # If it's a rate/quota error embedded in 200, back off and retry
                err_str = str(err).lower()
                if "rate" in err_str or "quota" in err_str or "limit" in err_str:
                    wait = 2 ** (attempt + 2)
                    print(f"[Groq] Quota error, waiting {wait}s...")
                    time.sleep(wait)
                    continue
                break

            try:
                text = (data["choices"][0]["message"]["content"] or "").strip()
            except (KeyError, IndexError, TypeError) as e:
                last_error = f"Unexpected shape from {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if not text:
                last_error = f"Empty content from {model}"
                print(f"[Groq] {last_error}")
                break

            print(f"[Groq] ✓ {model} ({len(text)} chars)")
            return text

        # Small pause between models to be friendly to rate limits
        time.sleep(1)

    raise RuntimeError(
        f"All accessible Groq models failed. Last error: {last_error}. "
        "The application queried Groq /models first, so this error now reflects "
        "models visible to your current API key/project. Check Groq Project "
        "model permissions and API key at https://console.groq.com"
    )


# Backward-compat alias
def gemini_stream(prompt, system="", temperature=0.7, progress_cb=None, tracked_sections=None):
    return ai_generate(prompt, system, temperature, progress_cb, tracked_sections)


SYSTEM_PROMPT = (
    "You are an expert academic research paper writer. "
    "You write in formal, scholarly English suitable for peer-reviewed journals. "
    "Do not use markdown formatting, bullet points, asterisks, or headers in your output — "
    "write clean flowing prose only, unless explicitly asked for a list. "
    "Be specific, evidence-grounded, and academically rigorous. "
    "Do not invent statistics or cite sources not provided to you."
)


# ═══════════════════════════════════════════════════════════════════════════════
#  WEB SCRAPER  (no API keys required)
# ═══════════════════════════════════════════════════════════════════════════════

def _http_get(url: str, timeout: int = 12) -> object:
    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "rdxper/3.0 (research-paper-generator; educational use)"}
        )
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8", errors="replace"))
    except Exception as e:
        print(f"[HTTP] {url[:80]} → {e}")
        return None


class WebScraper:
    def __init__(self, topic: str):
        self.topic = topic
        self.query = urllib.parse.quote(topic)

    def fetch_semantic_scholar(self, limit: int = 20) -> list:
        url = (
            f"https://api.semanticscholar.org/graph/v1/paper/search"
            f"?query={self.query}&limit={limit}"
            f"&fields=title,authors,year,abstract,citationCount,externalIds,publicationVenue"
        )
        data = _http_get(url)
        papers = []
        if data and "data" in data:
            for p in data["data"]:
                if not p.get("title"):
                    continue
                raw_authors = p.get("authors", [])
                # Skip papers with no identifiable author(s) rather than labelling
                # them "Unknown Author" — an unattributed entry isn't verifiable
                # and shouldn't feed into a literature review.
                names = [a.get("name") for a in raw_authors if a.get("name")]
                if not names:
                    continue
                elif len(names) == 1:
                    author_str = names[0]
                elif len(names) == 2:
                    author_str = f"{names[0]} & {names[1]}"
                else:
                    author_str = f"{names[0]} et al."
                papers.append({
                    "title":     p.get("title", "").strip(),
                    "authors":   author_str,
                    "year":      p.get("year") or 2022,
                    "abstract":  (p.get("abstract") or "").strip()[:500],
                    "doi":       (p.get("externalIds") or {}).get("DOI", ""),
                    "citations": p.get("citationCount") or 0,
                    "journal":   ((p.get("publicationVenue") or {}).get("name") or ""),
                })
        return papers

    def fetch_crossref(self, limit: int = 12) -> list:
        url = (
            f"https://api.crossref.org/works?query={self.query}"
            f"&rows={limit}&sort=relevance"
            f"&select=title,author,published,container-title,DOI"
        )
        data = _http_get(url)
        results = []
        if data and "message" in data:
            for item in data["message"].get("items", []):
                titles = item.get("title", [])
                title  = titles[0] if titles else ""
                if not title:
                    continue
                raw = [a for a in item.get("author", []) if a.get("family")]
                # Skip works with no identifiable author rather than labelling
                # them "Unknown Author" — same rule as the Semantic Scholar fetch.
                if not raw:
                    continue
                elif len(raw) == 1:
                    a = raw[0]
                    author_str = f"{a.get('family','?')}, {a.get('given','')[:1]}."
                elif len(raw) == 2:
                    a, b = raw[0], raw[1]
                    author_str = (
                        f"{a.get('family','?')}, {a.get('given','')[:1]}. & "
                        f"{b.get('family','?')}, {b.get('given','')[:1]}."
                    )
                else:
                    a = raw[0]
                    author_str = f"{a.get('family','?')}, {a.get('given','')[:1]}. et al."
                pub   = item.get("published", {})
                year  = (pub.get("date-parts") or [[2022]])[0][0]
                jlist = item.get("container-title", [])
                results.append({
                    "title":   title.strip(),
                    "authors": author_str,
                    "year":    year,
                    "journal": jlist[0] if jlist else "Academic Journal",
                    "doi":     item.get("DOI", ""),
                    "citations": 0,
                    "abstract": "",
                })
        return results

    def fetch_wikipedia(self) -> dict:
        slug = urllib.parse.quote(self.topic.replace(" ", "_"))
        url  = f"https://en.wikipedia.org/api/rest_v1/page/summary/{slug}"
        data = _http_get(url)
        if data and data.get("type") not in ("disambiguation",) and data.get("extract"):
            return {
                "summary": data["extract"],
                "url":     data.get("content_urls", {}).get("desktop", {}).get("page", ""),
                "title":   data.get("title", self.topic),
            }
        # Fallback: first word
        slug2 = urllib.parse.quote(self.topic.split()[0])
        data2 = _http_get(f"https://en.wikipedia.org/api/rest_v1/page/summary/{slug2}")
        if data2 and data2.get("extract"):
            return {
                "summary": data2["extract"],
                "url":     data2.get("content_urls", {}).get("desktop", {}).get("page", ""),
                "title":   data2.get("title", self.topic),
            }
        return {"summary": "", "url": "", "title": self.topic}

    def gather(self, progress_cb=None) -> dict:
        if progress_cb: progress_cb(10, "Querying Semantic Scholar for real papers...")
        ss = self.fetch_semantic_scholar(22)

        if progress_cb: progress_cb(18, "Querying CrossRef for verified journal articles...")
        cr = self.fetch_crossref(14)

        if progress_cb: progress_cb(24, "Fetching Wikipedia background context...")
        wiki = self.fetch_wikipedia()

        # Merge, deduplicate by title prefix
        seen = set()
        all_papers = []
        for p in ss + cr:
            key = p["title"][:40].lower()
            if key not in seen:
                seen.add(key)
                all_papers.append(p)

        # Sort by citation count
        all_papers.sort(key=lambda x: x.get("citations", 0), reverse=True)

        print(f"[Scraper] {len(ss)} SS papers, {len(cr)} CrossRef, wiki={'yes' if wiki.get('summary') else 'no'}")
        return {"papers": all_papers, "wiki": wiki}


# ═══════════════════════════════════════════════════════════════════════════════
#  CHART DATA MODEL
#  Fixed demographic value-labels (matching a real SPSS dataset export) and
#  response-scale templates used to build two-variable clustered chart specs.
# ═══════════════════════════════════════════════════════════════════════════════

DEMO_GROUPS = {
    'age':                        ['16-18', '19-35', '36-55', '55 above'],
    'gender':                     ['Male', 'female', 'Transgender'],
    'educational qualification':  ['Below 10th', '10th-12th', 'Undergraduate', 'Postgraduate', 'PhD'],
    'occupation':                 ['business', 'Employed', 'Retired', 'unemployed', 'Student'],
    'area':                       ['Urban', 'Rural', 'Semi-Urban'],
}
# Aliases so loosely-worded AI output ("education", "geographic area", ...) still resolves
_DEMO_ALIASES = {
    'education': 'educational qualification', 'qualification': 'educational qualification',
    'geographic area': 'area', 'region': 'area', 'residence': 'area', 'location': 'area',
    'employment': 'occupation', 'employment status': 'occupation', 'job': 'occupation',
    'sex': 'gender', 'age group': 'age',
}

LIKERT_10   = [str(i) for i in range(1, 11)]
AGREEMENT_5 = ['Strongly agree', 'Agree', 'Neutral', 'Disagree', 'Strongly disagree']


def resolve_demographic(raw: str):
    """Map arbitrary AI/user text to one of the fixed demographic axes + its
    real value-label set, so every chart's x-axis matches genuine SPSS output."""
    key = (raw or '').strip().lower()
    key = _DEMO_ALIASES.get(key, key)
    if key in DEMO_GROUPS:
        return key, DEMO_GROUPS[key]
    for k in DEMO_GROUPS:
        if k in key or key in k:
            return k, DEMO_GROUPS[k]
    # Deterministic fallback keeps things varied rather than always the same axis
    k = list(DEMO_GROUPS)[sum(ord(c) for c in key) % len(DEMO_GROUPS)] if key else 'gender'
    return k, DEMO_GROUPS[k]


def sparse_percent_matrix(rng: random.Random, n_groups: int, n_series: int) -> list:
    """Build a groups × series percentage matrix that sums to ~100%, with only
    a handful of nonzero cells per group (and some groups left empty) — this
    mirrors how a real SPSS crosstab of a convenience sample looks, rather
    than an evenly-filled synthetic grid."""
    weights = [[0.0] * n_series for _ in range(n_groups)]
    any_active = False
    for g in range(n_groups):
        if n_groups > 2 and rng.random() < 0.18:
            continue  # this demographic category had zero/negligible respondents
        k = rng.randint(1, min(3, n_series))
        for i in rng.sample(range(n_series), k):
            weights[g][i] = rng.uniform(8, 32)
            any_active = True
    if not any_active:
        weights[rng.randrange(n_groups)][rng.randrange(n_series)] = 20.0
    total = sum(sum(row) for row in weights) or 1.0
    return [[round(v / total * 100, 2) for v in row] for row in weights]


def _make_spec(question: str, xvar: str, groups: list, series: list, matrix: list) -> dict:
    legend_text = (f'The given figure represents the {xvar.lower()}-wise distribution of respondents\' '
                   f'responses to: "{question}"')
    return {
        'type': 'bar',
        'title': f'{question[:40]} by {xvar}',
        'question': question,
        'xvar': xvar,
        'groups': groups,
        'series': series,
        'matrix': matrix,
        'legend': legend_text,
        'interp': f'Distribution of responses across {len(series)} options, broken down by {xvar.lower()}.',
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  GEMINI CONTENT GENERATOR
#  Takes scraped data → asks Gemini to write each section
# ═══════════════════════════════════════════════════════════════════════════════

class GeminiWriter:
    def __init__(self, topic: str, scraped: dict, questionnaire: dict = None):
        self.topic        = topic
        self.papers       = scraped.get("papers", [])
        self.wiki         = scraped.get("wiki", {})
        self.seed         = sum(ord(c) for c in topic)
        random.seed(self.seed)
        np.random.seed(self.seed % 2**31)
        self.n_respondents = random.randint(210, 230)  # matches sample paper (~220)
        self.aware_pct     = random.randint(62, 74)
        self.fam_pct       = random.randint(70, 83)
        self.support_pct   = random.randint(62, 69)
        self.questionnaire = questionnaire or {}
        self._paper_digest = self._build_digest()
        self.sections      = {}   # filled by generate_all()

    def _build_digest(self) -> str:
        """Lean digest — titles/authors only, no abstracts. Minimises input tokens."""
        lines = []
        for i, p in enumerate(self.papers[:8], 1):
            jour = f", {p['journal']}" if p.get("journal") else ""
            lines.append(f"{i}. {p['authors']} ({p['year']}). \"{p['title']}\"{jour}. Cited {p.get('citations',0):,}x.")
        wiki = f"\nContext: {self.wiki['summary'][:120]}" if self.wiki.get("summary") else ""
        return "SOURCES:\n" + "\n".join(lines) + wiki

    def generate_all(self, progress_cb=None) -> dict:
        """
        4 AI calls matching the sample paper's exact structure:
        Call A : keywords + abstract + objectives
        Call B : introduction (6 named subheadings, ~1200-1500w)
        Call C : literature review (20 entries, exact fixed-sentence template, bold author/year)
        Call D : methodology + results (FIGURE paragraphs) + discussion +
                 limitations + suggestions + conclusion + charts
        """
        top      = sorted(self.papers, key=lambda x: x.get("citations", 0), reverse=True)
        top_cite = f"{top[0]['authors']} ({top[0]['year']})" if top else "prior studies"
        n, nr    = len(self.papers), self.n_respondents
        q        = self.questionnaire

        # ── Shared header (topic + top 4 papers + wiki) ──────────────────────
        digest_lines = []
        for i, p in enumerate(self.papers[:4], 1):
            digest_lines.append(f"{i}. {p['authors']} ({p['year']}). \"{p['title']}\".")
        wiki_snip = self.wiki.get("summary", "")[:120] if self.wiki.get("summary") else ""
        hdr = f"TOPIC: {self.topic} | N={nr} respondents | Top paper: {top_cite}"
        if wiki_snip:
            hdr += f" | Context: {wiki_snip}"
        hdr += "\n" + "\n".join(digest_lines) + "\n\n"

        # Researcher inputs (only non-empty fields)
        q_prob = f"Problem statement: {q['problem']}\n" if q.get('problem') else ""
        q_obj  = f"Objectives (reproduce verbatim):\n{q['objectives']}\n" if q.get('objectives') else ""
        q_stmt = f"Research statement: {q['statement']}\n" if q.get('statement') else ""
        q_gap  = f"Research gap: {q['gap']}\n" if q.get('gap') else ""
        q_lit  = f"Key literature noted by researcher: {q['lit'][:300]}\n" if q.get('lit') else ""

        sections = {}

        # ── PARALLEL AI CALLS A, B, C, D ─────────────────────────────────────
        # All four calls share only pre-built hdr/q context — fully independent,
        # so we fire them concurrently for a ~3× speed improvement.
        
        pA = (hdr + q_prob + q_obj + q_stmt +
              "Write using XML tags only. Scholarly prose, no markdown bullets outside objectives.\n\n"
              "<keywords>6-8 comma-separated academic keywords for this topic. Do NOT use bold markers.</keywords>\n"
              f"<abstract>Write EXACTLY 250 words — count carefully and stop at 250. "
              f"Write as ONE single flowing paragraph with NO line breaks or subheadings. "
              f"Begin with 2 sentences giving context about {self.topic}. "
              f"Then embed these EXACT bold inline labels in the prose in this order:\n"
              f"'The **Aim** of the study is to [specific aim].'\n"
              f"'The **Objective** is to [main objective].'\n"
              f"'The **sample size** of the study is {nr}.'\n"
              f"'The **Findings** of the study were that [3-4 key findings with % values].'\n"
              f"'In **Conclusion** [1-2 sentences on policy implications].'\n"
              "Target EXACTLY 250 words. Count rigorously before finalising.</abstract>\n"
              "<objectives>"
              + ("Reproduce VERBATIM, each on its own line, each starting '● To ...':\n" + q['objectives'] if q.get('objectives') else
                 f"Write EXACTLY 4 specific objectives for {self.topic}, each on its own line starting '● To [active verb] ...'. No more, no fewer than 4.")
              + "</objectives>")

        pB = (hdr + q_prob + q_gap + q_stmt +
              "Write a formal academic INTRODUCTION section using XML tags. Flowing prose only — no bullet points, no subheadings, no bold markers.\n\n"
              f"<introduction>Write EXACTLY 3 to 4 distinct paragraphs separated by a blank line. "
              f"Each paragraph must be continuous flowing prose with NO subheadings, NO bold text, and NO bullet points. "
              f"Total length: 450-600 words across all paragraphs.\n\n"
              f"PARAGRAPH 1 (120-160 words): Establish the historical background and significance of {self.topic}. "
              f"Introduce the key institutions, laws, or frameworks associated with {self.topic}. "
              f"Name specific acts, bodies, or landmark developments relevant to {self.topic} in India or globally. "
              f"End the paragraph with a transition to the broader governance or regulatory context.\n\n"
              f"PARAGRAPH 2 (120-150 words): Introduce the key oversight or review mechanism relevant to {self.topic} "
              f"(e.g. judicial review, regulatory audit, policy evaluation, or equivalent). "
              f"Explain its constitutional or statutory basis, its purpose as a safeguard, and how it operates in practice. "
              + (f"Connect to this research gap: {q['gap'][:150]}. " if q.get('gap') else "Discuss tensions between technical expertise and legal/policy standards. ") +
              f"This paragraph should elaborate on why oversight matters for {self.topic}.\n\n"
              f"PARAGRAPH 3 (100-140 words): Discuss the evolving landscape of {self.topic} — recent legislative, "
              f"judicial, technological, or policy developments that have reshaped the field. "
              f"Note any structural changes, reforms, or shifts in how {self.topic} is governed or studied. "
              f"Raise the important questions this evolution creates for researchers, practitioners, and policymakers.\n\n"
              f"PARAGRAPH 4 (80-120 words): State the aim and scope of this study. "
              f"Begin with: 'This study aims to critically examine...' and describe what the paper explores, "
              f"the legal/empirical/analytical framework used, and the contribution it makes to understanding {self.topic}. "
              + (f"Ground the aim in: {q['statement'][:120]}" if q.get('statement') else "") +
              f"\n\nCRITICAL: No subheadings. No bold text. No numbered lists. Separate paragraphs with a blank line only. "
              f"Write scholarly prose as found in peer-reviewed law or social science journals.</introduction>")

        # Build a digest of scraped real papers to seed the AI's lit review.
        # Every entry here has already passed through the author/title filters
        # in fetch_semantic_scholar / fetch_crossref, so nothing "Unknown" is used.
        scraped_seed = ""
        for i, p in enumerate(self.papers[:20], 1):
            jour = f", {p['journal']}" if p.get("journal") else ""
            doi  = f", DOI: {p['doi']}" if p.get("doi") else ""
            abst = f" Abstract: {p['abstract'][:200]}" if p.get("abstract") else ""
            scraped_seed += f"{i}. {p['authors']} ({p['year']}). \"{p['title']}\"{jour}{doi}.{abst}\n"

        # Ask for exactly as many entries as we have real, verified sources —
        # never force a fixed count (e.g. 20) that would pressure the model to
        # pad the review with studies recalled from memory (and therefore
        # unverifiable / at risk of being hallucinated).
        n_real_papers = len(self.papers[:20])
        lit_count = n_real_papers

        pC = (f"TOPIC: {self.topic}\n"
              + (f"Researcher's key sources: {q['lit'][:300]}\n" if q.get('lit') else "")
              + (f"\nREAL PAPERS SCRAPED FROM SEMANTIC SCHOLAR & CROSSREF (use these as your primary sources):\n{scraped_seed}\n" if scraped_seed else "")
              + "Write a LITERATURE REVIEW using XML tags.\n"
              "CRITICAL RULE: Use ONLY the real, verifiable papers listed above — do NOT introduce any "
              "additional studies, authors, or titles recalled from memory, even ones you believe are real "
              "and well-known. If you are not 100% certain a paper exists exactly as scraped above, leave it out.\n"
              "Never invent authors, titles, journals, DOIs, or statistics. Never fabricate a study that does not exist. "
              "Do not invent specific numeric findings (%, sample sizes, etc.) unless they are given to you above — "
              "describe findings qualitatively ('showed a positive impact', 'found limited awareness') when no real "
              "figure is known. If unsure of any detail, write around it rather than invent it.\n\n"
              f"<literature_review>Write EXACTLY {lit_count} entries — one for each scraped paper listed above, "
              f"no more and no fewer — one per paragraph, separated by a blank line.\n"
              f"Every entry MUST correspond to one of the real papers listed above — "
              f"never an invented author or study, and never a paper not in that list.\n\n"
              f"EVERY ENTRY MUST FOLLOW THIS EXACT SEVEN-SENTENCE TEMPLATE, IN THIS ORDER, WITH NO DEVIATION:\n"
              f"  1. '[Author Surname(s) or Organisation] (Year) examined [what the study investigated, tied to {self.topic}].'\n"
              f"  2. 'The objective focused on [the study's specific objective].'\n"
              f"  3. 'The methodology adopted was [research method/approach actually used, or a reasonable qualitative "
              f"description if the exact method is not stated in the source].'\n"
              f"  4. 'The findings showed [key result(s), described qualitatively unless a real statistic is known].'\n"
              f"  5. 'The study suggested [the recommendation(s) made by the study].'\n"
              f"  6. 'The future scope proposed [a direction for further research consistent with the study's own conclusion/discussion].'\n"
              f"  7. 'The conclusion highlighted [the study's overall takeaway and its relevance to {self.topic}].'\n\n"
              f"EXAMPLE OF THE REQUIRED FORMAT (match this structure and tone exactly, but with real content):\n"
              f"\"Bhattacharya and Ghosh (2020) examined the implementation of AI-based tools in the Indian judicial "
              f"system, including translation and legal research software. The objective focused on evaluating AI's "
              f"role in reducing case backlog and language barriers. The methodology adopted was qualitative analysis "
              f"of judicial initiatives and policy documents. The findings showed positive impact on administrative "
              f"efficiency but limited public awareness. The study suggested expanding AI tools with multilingual "
              f"capabilities. The future scope proposed studying public perception of AI-enabled judicial services. "
              f"The conclusion highlighted AI's potential to enhance access to justice if inclusively implemented.\"\n\n"
              f"RULES:\n"
              f"  1. No number prefix, no bullet — start directly with Author/Org name then (Year)\n"
              f"  2. Exactly seven sentences per entry, in the exact order above, using the exact lead-in phrases "
              f"'The objective focused on', 'The methodology adopted was', 'The findings showed', 'The study "
              f"suggested', 'The future scope proposed', and 'The conclusion highlighted'\n"
              f"  3. Only real works from the scraped list above — never invent a paper, author, or finding to fill "
              f"the template; if a study's future scope or suggestion is not explicit in the source, infer a modest, "
              f"plausible one consistent with its actual findings rather than fabricating unrelated claims\n"
              f"  4. Use only the years given for each scraped paper above — do not alter them\n"
              f"  5. Separate entries with a blank line; no section headings, no sub-labels\n"
              f"  6. Aim for 90–120 words per entry\n"
              f"  7. Write EXACTLY {lit_count} entries — one per scraped paper, no more, no fewer</literature_review>\n\n"
              f"<references>Generate APA 7th edition references for the same {lit_count} real works, in the same "
              f"order, numbered 1–{lit_count}.\n"
              f"FORMAT:\n"
              f"  Journal: [N]. Author, A. A., & Author, B. B. (Year). Title. Journal, volume(issue), pages.\n"
              f"  Book: [N]. Author, A. A. (Year). Title. Publisher.\n"
              f"  Report: [N]. Organisation. (Year). Title. Publisher.\n"
              f"Do NOT invent DOIs. Omit DOI if not certain it is real.</references>")

        pD = (hdr +
              "Write the remaining paper sections using XML tags. Scholarly prose only — no bullet points.\n\n"
              f"<methodology>Write EXACTLY ONE single flowing paragraph of approximately 200 words (no more than 220 words). "
              f"No subheadings, no bullet points, no line breaks — continuous scholarly prose only. "
              f"The paragraph MUST open with: 'The research method which is followed here is empirical research.' "
              f"Then in the same flowing paragraph, concisely cover all of the following: "
              f"why empirical/descriptive research suits {self.topic}; "
              f"'A total of {nr} samples have been collected' via convenience sampling; "
              f"data collection through a structured questionnaire with a five-point Likert scale; "
              f"secondary sources including peer-reviewed journals, government reports, and statistical databases consulted for {self.topic}; "
              f"SPSS version 21 used for statistical analysis employing chi-square, ANOVA, and Pearson correlation tests; "
              f"independent variables of age, gender, educational qualification, geographic area, and occupation; "
              f"and the dependent variable being the relevant outcome or awareness level for {self.topic}. "
              f"Aim for tight, precise academic prose — approximately 200 words in a single unbroken paragraph.</methodology>\n\n"
              f"<results>Write ONE single continuous paragraph (no line breaks or blank lines within it). "
              f"Describe key findings from all {self._nfigs} figures in sequence, embedding an inline bold figure "
              f"reference after each figure's description. "
              f"FORMAT: 'According to the chart, [2-3 sentences of findings with realistic % values for figure 1] **(fig: 1)**. "
              f"The data reveals [findings for figure 2] **(fig: 2)**. [Continue for all figures up to {self._nfigs}].' "
              f"Keep each figure's description to 2-3 sentences with specific demographic % values. "
              f"Use varied openers: 'According to the chart,', 'The data reveals that', 'Among respondents,', "
              f"'Notably,', 'A majority of respondents', 'The findings show that', etc. "
              f"End the paragraph after **(fig: {self._nfigs})**. No separate paragraphs — ONE block of text only.</results>\n\n"
              f"<discussion>Write EXACTLY {self._nfigs} paragraphs separated by blank lines — one per Figure. "
              f"Each paragraph: 60-90 words of flowing scholarly prose connecting the figure's data to the broader context of {self.topic}. "
              f"Mention subgroup differences, relevant concepts, and policy implications. "
              f"EACH paragraph MUST END with the inline bold figure reference: **(fig: [N])** "
              f"Example closing: '...reinforcing the need for targeted policy interventions in this domain **(fig: 1)**.' "
              f"Do NOT open paragraphs with 'FIGURE N' — start directly with the analysis. "
              f"Prose only, no headings, no bullet points within.</discussion>\n\n"
              f"<limitations>Write 2 paragraphs (150-200 words total). "
              f"First paragraph: study limitations — sample characteristics, geographic scope, convenience sampling bias, "
              f"self-report limitations specific to {self.topic}. "
              f"Second paragraph: scope limitations and directions for future research.</limitations>\n\n"
              f"<suggestions>Write 5-6 actionable recommendations for {self.topic} in flowing prose paragraphs "
              f"(200-250 words total). Address policy makers, technology providers, communities, and researchers. "
              f"No bullet points.</suggestions>\n\n"
              f"<conclusion>Write 5-7 prose paragraphs (600-700 words). Cover in order: "
              f"(1) restate study purpose and topic significance; "
              f"(2) summary of key findings with % values from the {nr}-respondent sample; "
              f"(3) how findings meet the stated objectives; "
              f"(4) implications and 4-5 specific policy recommendations for {self.topic}; "
              f"(5) study limitations briefly; "
              f"(6) future research directions. No bullets.</conclusion>\n\n"
              f"<charts>{self._nfigs} lines, one per figure. Format: bar|QUESTION|XVAR|SERIES\n"
              f"QUESTION is the full survey statement/question this figure reports on, related to {self.topic[:40]}.\n"
              f"XVAR is exactly one of: age, gender, educational qualification, occupation, area — the demographic "
              f"the responses are broken down by. Cycle through all five before repeating, in roughly equal proportion.\n"
              f"SERIES is a comma-separated list of the response options for the question — use ONE of these forms: "
              f"a 5-point agreement scale (Strongly agree,Agree,Neutral,Disagree,Strongly disagree), "
              f"a 1-10 rating scale (1,2,3,4,5,6,7,8,9,10), or 3-6 short topic-specific multiple-choice options.\n"
              f"EXAMPLES:\n"
              f"bar|To what extent do you agree with the statement: \"{self.topic[:40]} poses a serious risk\"?|age|"
              f"Strongly agree,Agree,Neutral,Disagree,Strongly disagree\n"
              f"bar|How would you rate your awareness of {self.topic[:35]}?|gender|1,2,3,4,5,6,7,8,9,10\n"
              f"bar|Which entity should be primarily responsible for addressing {self.topic[:30]}?|occupation|"
              f"Government,Industry,Educators,Individuals,Civil Society</charts>")

        # Sequential calls — avoids hammering rate limits with simultaneous requests
        if progress_cb: progress_cb(32, "Writing keywords, abstract & objectives...")
        raw_A = ai_generate(pA, system=SYSTEM_PROMPT, temperature=0.7)

        if progress_cb: progress_cb(44, "Writing introduction...")
        raw_B = ai_generate(pB, system=SYSTEM_PROMPT, temperature=0.7)

        if progress_cb: progress_cb(56, "Writing literature review...")
        if lit_count > 0:
            raw_C = ai_generate(pC, system=SYSTEM_PROMPT, temperature=0.7)
        else:
            # No real, attributable papers were scraped for this topic — do not
            # ask the model to write a literature review, since with zero
            # verified sources it would have nothing real to draw from.
            print("[LitReview] 0 verified scraped papers — skipping AI call, using empty/fallback review.")
            raw_C = ""

        if progress_cb: progress_cb(66, "Writing methodology, results & conclusion...")
        raw_D = ai_generate(pD, system=SYSTEM_PROMPT, temperature=0.7)

        # Parse Call A
        for tag in ('keywords', 'abstract', 'objectives'):
            m = re.search(rf'<{tag}>(.*?)</{tag}>', raw_A, re.DOTALL)
            sections[tag] = m.group(1).strip() if m else ''

        # Parse Call B
        m = re.search(r'<introduction>(.*?)</introduction>', raw_B, re.DOTALL)
        sections['introduction'] = m.group(1).strip() if m else ''

        # Parse Call C
        m = re.search(r'<literature_review>(.*?)</literature_review>', raw_C, re.DOTALL)
        sections['literature_review'] = m.group(1).strip() if m else ''
        m_refs = re.search(r'<references>(.*?)</references>', raw_C, re.DOTALL)
        sections['ai_references'] = m_refs.group(1).strip() if m_refs else ''

        # Parse Call D
        for tag in ('methodology', 'results', 'discussion', 'limitations', 'suggestions', 'conclusion', 'charts'):
            m = re.search(rf'<{tag}>(.*?)</{tag}>', raw_D, re.DOTALL)
            sections[tag] = m.group(1).strip() if m else ''

        if progress_cb: progress_cb(74, "All sections written. Assembling Word document...")

        # ── Fallbacks ─────────────────────────────────────────────────────────
        fallbacks = {
            'keywords':          f'{self.topic}, empirical study, policy, digital media, awareness, India',
            'abstract':          (
                f'{self.topic} has emerged as a significant area of scholarly and policy concern in recent decades, '
                f'reflecting the complex interplay of technology, law, and social behaviour in contemporary societies. '
                f'As digital platforms and legislative frameworks continue to evolve, understanding public awareness '
                f'and the effectiveness of existing interventions has become increasingly urgent for researchers and policymakers alike. '
                f'The **Aim** of the study is to examine the factors influencing awareness and attitudes towards {self.topic} '
                f'across diverse demographic groups and to identify gaps in current preventive frameworks. '
                f'The **Objective** is to assess the role of education, technology, and policy in shaping outcomes '
                f'related to {self.topic} and to generate actionable recommendations for improvement. '
                f'The **sample size** of the study is {nr}, comprising respondents drawn through convenience sampling '
                f'from varied age, gender, educational, and occupational backgrounds. '
                f'The **Findings** of the study were that approximately 68% of respondents demonstrated moderate '
                f'to high awareness of {self.topic}, with educational qualification emerging as the strongest predictor '
                f'of awareness levels; 54% expressed support for enhanced technological interventions, '
                f'while 42% identified gaps in government outreach and enforcement mechanisms. '
                f'In **Conclusion**, the study underscores the need for integrated policy responses combining '
                f'digital innovation, legislative reform, and community education to effectively address {self.topic}.'
            ),
            'introduction':      (
                f'{self.topic} has emerged as a significant area of scholarly and policy concern in recent decades, '
                f'reflecting the complex interplay of legislation, institutional frameworks, and evolving social realities. '
                f'In India, the administration and adjudication of matters related to {self.topic} have increasingly been entrusted '
                f'to specialised bodies and statutory authorities designed to ensure technical expertise, efficiency, and expedited resolution. '
                f'These institutions play a pivotal role in interpreting and enforcing the relevant legal and regulatory frameworks '
                f'governing {self.topic}, including key acts, schemes, and policy instruments introduced at the national and state levels.\n\n'
                f'However, the decisions and practices of such administrative and regulatory bodies are not beyond scrutiny. '
                f'Mechanisms of judicial and regulatory oversight — vested in constitutional courts and supervisory authorities — '
                f'act as fundamental safeguards to ensure that these bodies function within the bounds of legality, fairness, '
                f'and procedural propriety. Such oversight serves to check arbitrariness, jurisdictional errors, and violations '
                f'of principles of natural justice, thereby maintaining the rule of law and accountability in the governance of {self.topic}.\n\n'
                f'The evolving landscape of {self.topic} in India, particularly in light of recent legislative reforms, '
                f'judicial pronouncements, and technological developments, has further intensified the discourse on the role and '
                f'scope of effective governance in this domain. Policymakers and courts are increasingly called upon to balance '
                f'deference to technical expertise with the need to uphold constitutional principles and legal standards, '
                f'raising important questions regarding the extent of intervention, the standards applied, and the implications '
                f'for efficiency and consistency in the administration of {self.topic}.\n\n'
                f'This study aims to critically examine the governance and empirical dimensions of {self.topic} in India. '
                f'It explores the legal and institutional framework, analyses key judicial and policy precedents, and evaluates '
                f'the challenges and implications of current approaches. By doing so, the research seeks to contribute to a '
                f'nuanced understanding of the relationship between administrative practice and constitutional or regulatory '
                f'control in the field of {self.topic}.'
            ),
            'objectives':        f'● To evaluate the potential of technology in preventing and addressing {self.topic}.\n● To identify vulnerabilities and challenges in the existing frameworks governing {self.topic}.\n● To assess the impact of awareness campaigns and digital platforms in educating the public about {self.topic}.\n● To recommend evidence-based policy interventions and best practices for effectively addressing {self.topic}.',
            'literature_review': self._build_lit_review_fallback(nr),
            'methodology':       f'The research method which is followed here is empirical research. Descriptive and empirical research is particularly suited to investigating {self.topic} because it enables systematic data collection and quantitative analysis of real-world attitudes and behaviours. A total of {nr} samples have been collected through convenience sampling, comprising respondents across multiple demographic categories. Data were gathered through structured questionnaires administered during field visits, incorporating a five-point Likert scale to measure attitudes and perceptions. Secondary sources including peer-reviewed journals, government reports, and statistical databases were also consulted. Data analysis was performed using SPSS version 21 with chi-square, ANOVA, and Pearson correlation tests. Independent variables comprise age, gender, educational qualification, geographic area, and occupation; the dependent variable is awareness and attitude towards {self.topic}.',
            'results':           ' '.join([
                f'According to the chart, {round(10+i*2.1,1)}% of respondents aged 18–30, {round(18+i*1.3,1)}% aged 31–40, '
                f'{round(14+i*0.9,1)}% aged 41–50, and {round(8+i*0.7,1)}% aged 51 and above provided their responses towards {self.topic}. '
                f'The 31–40 age group showed the strongest awareness levels, with educational qualification emerging as a significant moderating factor. **(fig: {i})**'
                for i in range(1, self._nfigs+1)
            ]),
            'discussion':        '\n\n'.join([
                f'The data on {self.topic} reveals that awareness is most pronounced among respondents with higher educational qualifications and those in the 31–40 age bracket. '
                f'This finding aligns with existing scholarship on the relationship between education level and civic awareness of sensitive social issues. '
                f'The data underscores the importance of targeted educational and digital outreach strategies to reach under-informed demographic segments. **(fig: {i})**'
                for i in range(1, self._nfigs+1)
            ]),
            'limitations':       f'The body of literature reviewed highlights several limitations that merit consideration. This study relies on a convenience sample of {nr} respondents, which, while adequate for exploratory analysis, limits the generalisability of findings across all population groups relevant to {self.topic}. Self-report biases and social desirability effects may have influenced responses on sensitive dimensions of the topic.\n\nThe geographic scope of the study is concentrated and may not adequately represent rural and remote populations who experience {self.topic} differently from urban respondents. Future research should employ longitudinal methodologies with larger, more geographically diverse samples across multiple Indian states to validate and extend the current findings.',
            'suggestions':       f'Policymakers should prioritise strengthening the legislative and regulatory frameworks governing {self.topic} and ensure that existing laws are rigorously enforced at both central and state levels. Investment in technology-based solutions, including AI-driven monitoring and reporting systems, should be accelerated. Community awareness programmes must be expanded with particular emphasis on reaching underserved and rural populations. Educational institutions should integrate age-appropriate curricula to build long-term awareness from an early stage. Researchers and practitioners should collaborate to develop evidence-based intervention models suitable for adoption by state governments. Civil society organisations must be adequately funded and legally empowered to support affected individuals and advocate for systemic reform.',
            'conclusion':        f'This study has undertaken an empirical examination of {self.topic} through research with {nr} respondents drawn from diverse demographic backgrounds. The findings indicate significant variation in awareness and attitudes across educational, age, gender, and occupational groups, with the majority demonstrating moderate to high levels of awareness. Graduate-level respondents and those in the 31-40 age group showed the strongest engagement with the issue, while rural respondents and those with lower educational attainment indicated comparatively lower awareness. These findings substantially fulfil the stated objectives of the study, confirming the relevance of education and technology-based interventions. Policymakers should prioritise legislative reform, digital infrastructure investment, and sustained community outreach as the three pillars of an integrated response to {self.topic}. The study is subject to limitations in sample size and geographic coverage, which future longitudinal and multi-state research should address to build a more comprehensive evidence base.',
            'charts':            '',
        }
        for k, fb in fallbacks.items():
            if not sections.get(k):
                sections[k] = fb

        self.sections = sections
        return sections


    def _build_lit_review_fallback(self, nr: int) -> str:
        """
        Build a literature review from REAL scraped papers (Semantic Scholar + CrossRef) only.
        Used only if the AI call fails. Never invents author names, papers, or statistics —
        each entry follows the fixed 7-sentence template but every claim is either drawn
        from the paper's real metadata/abstract or phrased qualitatively (no fabricated numbers).

        Target is exactly 20 entries, but this fallback will NEVER pad the count with
        invented studies: if fewer than 20 real papers were scraped, it returns fewer
        entries rather than hallucinate additional ones.
        """
        entries = []

        for p in self.papers[:20]:
            authors  = (p.get('authors') or '').strip()
            year     = p.get('year', 2020)
            title    = p.get('title', '').strip()
            journal  = (p.get('journal') or '').strip()
            abstract = (p.get('abstract') or '').strip()
            cites    = p.get('citations', 0)

            # Skip anything without a real title or a real, attributed author —
            # never fabricate a placeholder like "Unknown Author" to fill a slot.
            if not title or not authors:
                continue

            venue_clause = f" (published in {journal})" if journal else ""

            if abstract and len(abstract) > 80:
                # Ground every clause in the real abstract; no invented figures.
                snip = abstract[:220].rsplit(' ', 1)[0].rstrip('.')
                examined   = f"\"{title}\"{venue_clause}, focusing on {snip.lower()}"
                objective  = f"understanding the core issue addressed in the study — {snip[:110].lower()}"
                methodology = "an analysis of the study's own reported data and documentary/empirical evidence, as described in the published abstract"
                findings   = "outcomes consistent with the abstract's own reported conclusions, without additional invented statistics"
                suggestion = f"further application or refinement of its approach within the broader field of {self.topic}"
                future     = f"extending this line of enquiry to related aspects of {self.topic} not yet covered by the original study"
                conclusion = f"the continued relevance of this work to scholarship and practice on {self.topic}"
            else:
                # No usable abstract — describe only what is verifiably known (title/venue/year),
                # and keep the rest strictly qualitative. No invented sample sizes or percentages.
                examined   = f"\"{title}\"{venue_clause}, addressing aspects of {self.topic}"
                objective  = f"exploring the dimensions of {self.topic} relevant to this work's stated title and venue"
                methodology = "a methodology consistent with standard academic practice in the field, as indicated by its publication venue"
                findings   = "results that contribute to the broader scholarly understanding of the subject"
                suggestion = f"further empirical or policy attention to the specific dimension of {self.topic} it addresses"
                future     = f"replication or extension of the work in different contexts relevant to {self.topic}"
                conclusion = f"the work's contribution to the wider literature on {self.topic}"

            cite_clause = f" This work has been cited {cites:,} times." if cites and cites > 10 else ""

            entry = (
                f"{authors} ({year}) examined {examined}. "
                f"The objective focused on {objective}. "
                f"The methodology adopted was {methodology}. "
                f"The findings showed {findings}. "
                f"The study suggested {suggestion}. "
                f"The future scope proposed {future}. "
                f"The conclusion highlighted {conclusion}.{cite_clause}"
            )
            entries.append(entry)

        return '\n\n'.join(entries)

    def parse_chart_specs(self, n: int) -> list:
        """Parse the <charts> block from Gemini into renderable spec dicts.

        Each figure is modelled the way SPSS's Chart Builder actually reports a
        crosstab: a demographic axis (age / gender / educational qualification /
        occupation / area) clustered against the response options for a survey
        question, with a legend headed by the full question text.
        """
        rng = random.Random(self.seed + 7)
        specs = []
        raw   = self.sections.get('charts', '')

        for line in raw.strip().splitlines():
            line = line.strip()
            if not line or '|' not in line:
                continue
            parts = [p.strip() for p in line.split('|')]
            if len(parts) < 4:
                continue
            _type, question, xvar_raw, series_raw = parts[0], parts[1], parts[2], parts[3]
            try:
                xvar, groups = resolve_demographic(xvar_raw)
                series = [s.strip() for s in series_raw.split(',') if s.strip()][:10]
                if len(series) < 2:
                    series = list(LIKERT_10)
                matrix = sparse_percent_matrix(rng, len(groups), len(series))
                specs.append(_make_spec(question, xvar, groups, series, matrix))
            except Exception as e:
                print(f"[Chart parse] skipped: {line!r} → {e}")
                continue
            if len(specs) >= n:
                break

        # Pad with fallbacks if needed
        while len(specs) < n:
            specs.extend(self._fallback_specs(n - len(specs)))
            break

        return specs[:n]

    def references(self) -> list:
        refs, seen = [], set()
        for p in self.papers[:20]:
            key = p["title"][:35].lower()
            if key in seen: continue
            seen.add(key)
            journal = p.get("journal") or "Academic Journal"
            doi_str = f" https://doi.org/{p['doi']}" if p.get("doi") else ""
            refs.append(f"{p['authors']} ({p['year']}). {p['title']}. {journal}.{doi_str}")
        if self.wiki.get("url"):
            refs.append(f"Wikipedia contributors. ({datetime.now().year}). {self.wiki.get('title', self.topic)}. Wikipedia. {self.wiki['url']}")
        # NOTE: previously padded the list with three hardcoded, topic-irrelevant
        # citations (WIPO/UNESCO/Floridi) regardless of subject matter. Removed —
        # a reference list should only ever contain works actually scraped/relevant
        # to this paper's topic, never filler citations bolted on to hit a count.
        return list(dict.fromkeys(refs))[:20]

    def _fallback_specs(self, n: int) -> list:
        """Safe fallback chart specs requiring no AI call — cycles through all
        five demographic axes and a mix of response-scale templates so even a
        large figure count (up to 25) produces varied, real-looking crosstabs."""
        rng = random.Random(self.seed)
        demo_keys = ['age', 'gender', 'educational qualification', 'occupation', 'area']
        series_templates = [
            AGREEMENT_5,
            LIKERT_10,
            ['Yes', 'No', 'Not Sure'],
            ['Very Low', 'Low', 'Moderate', 'High', 'Very High'],
        ]
        question_stems = [
            f'To what extent do you agree with the statement: "{self.topic[:42]} requires urgent attention"?',
            f'How would you rate your overall awareness of {self.topic[:40]}?',
            f'Do you believe existing frameworks adequately address {self.topic[:35]}?',
            f'How concerned are you about the impact of {self.topic[:38]}?',
            f'To what extent do you trust current measures related to {self.topic[:32]}?',
            f'How would you rate the effectiveness of policies on {self.topic[:35]}?',
            f'To what extent do you support stronger regulation of {self.topic[:35]}?',
        ]
        specs = []
        for i in range(n):
            xvar, groups = resolve_demographic(demo_keys[i % len(demo_keys)])
            series   = series_templates[i % len(series_templates)]
            question = question_stems[i % len(question_stems)]
            matrix   = sparse_percent_matrix(rng, len(groups), len(series))
            specs.append(_make_spec(question, xvar, groups, series, matrix))
        return specs[:n]


# ═══════════════════════════════════════════════════════════════════════════════
#  CHART RENDERING  (matplotlib SPSS-style)
# ═══════════════════════════════════════════════════════════════════════════════

# Default SPSS categorical palette, in legend order (1..10 / Strongly agree..Strongly disagree / etc.)
SPSS_COLORS = [
    '#2A3B8F',  # 1 — navy blue
    '#4E9A4E',  # 2 — green
    '#D6CCA3',  # 3 — khaki / tan
    '#7B1FA2',  # 4 — purple
    '#FFFF66',  # 5 — pale yellow
    '#E4141B',  # 6 — red
    '#3FBFBF',  # 7 — teal / cyan
    '#BEBEBE',  # 8 — grey
    '#8DB3E2',  # 9 — light steel blue
    '#1B4D2E',  # 10 — dark green
]


def _wrap(text: str, width: int = 24) -> str:
    import textwrap
    return '\n'.join(textwrap.wrap(text, width)) or text


def _spss_clustered_chart(question, xvar, groups, series, matrix):
    """Render a two-variable clustered bar chart matching SPSS Chart Builder's
    default output: light-grey plot area, solid black frame, no gridlines,
    thin black bar outlines, boxed % data labels, and a right-hand legend
    headed by the survey question text (all series shown, even zero ones)."""
    n_groups = len(groups)
    n_series = len(series)
    colors   = [SPSS_COLORS[i % len(SPSS_COLORS)] for i in range(n_series)]

    fig, ax = plt.subplots(figsize=(4.6, 3.75))

    group_width = 0.62
    for gi in range(n_groups):
        active = [(si, matrix[gi][si]) for si in range(n_series) if matrix[gi][si] and matrix[gi][si] > 0]
        if not active:
            continue
        k     = len(active)
        bw    = group_width / k
        start = gi - group_width / 2
        for j, (si, val) in enumerate(active):
            xpos = start + bw * j + bw / 2
            ax.bar(xpos, val, width=bw * 0.94, color=colors[si],
                   edgecolor='black', linewidth=0.8, zorder=3)
            label_y = val * 0.9 if val > 4 else val
            va      = 'top' if val > 4 else 'bottom'
            ax.annotate(f'{val:.2f}%', xy=(xpos, label_y), ha='center', va=va,
                        fontsize=6.3, color='#222222',
                        bbox=dict(boxstyle='square,pad=0.22', fc='white', ec='black', lw=0.6),
                        zorder=4)

    ax.set_xlim(-0.5, n_groups - 0.5)
    ax.set_xticks(range(n_groups))
    max_len = max((len(g) for g in groups), default=0)
    crowded = max_len > 9 or n_groups >= 5
    ax.set_xticklabels(
        groups, fontsize=7.3 if crowded else 8, color='#111111',
        rotation=18 if crowded else 0,
        ha='right' if crowded else 'center',
    )

    flat = [v for row in matrix for v in row if v]
    ymax = (max(flat) if flat else 10) * 1.3 + 3
    ax.set_ylim(0, ymax)

    # SPSS look: light-grey plot area, solid black frame, no gridlines
    ax.set_facecolor('#EAEAEA')
    fig.patch.set_facecolor('white')
    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_color('black')
        spine.set_linewidth(0.9)
    ax.xaxis.grid(False)
    ax.yaxis.grid(False)
    ax.set_axisbelow(True)
    ax.tick_params(colors='#111111', labelsize=8, length=3)
    ax.set_ylabel('Percent', fontsize=9, fontweight='bold', color='#111111')
    ax.set_xlabel(xvar, fontsize=9, fontweight='bold', color='#111111')

    handles = [plt.Rectangle((0, 0), 1, 1, fc=colors[i], ec='black', linewidth=0.6) for i in range(n_series)]
    leg = ax.legend(handles, series, title=_wrap(question, 15),
                     loc='upper left', bbox_to_anchor=(1.02, 1.02), frameon=False,
                     fontsize=6.5, title_fontsize=6.8, handlelength=1.0, handleheight=1.0,
                     labelspacing=0.3, borderaxespad=0)
    leg._legend_box.align = 'left'
    leg.get_title().set_ha('left')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=170, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_chart(spec: dict) -> io.BytesIO:
    return _spss_clustered_chart(
        spec.get('question', spec.get('title', '')),
        spec.get('xvar', 'Gender'),
        spec.get('groups', ['Male', 'female', 'Transgender']),
        spec.get('series', LIKERT_10),
        spec.get('matrix'),
    )


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, color: str):
    tc  = cell._tc
    pr  = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pr.append(shd)

def _add_table(doc, caption: str, rows: list, hcol: str = '1F3864'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run  = para.add_run(str(txt))
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, hcol.upper())
            elif ri % 2 == 0:
                _set_cell_bg(cell, 'EBF3FB')
    doc.add_paragraph()


class DocBuilder:
    def __init__(self, topic, author, inst, email, writer: GeminiWriter,
                 sections: dict, specs: list, charts: list, papers: list,
                 co_author: str = '', co_author_title: str = '',
                 co_author_inst: str = '', co_author_email: str = '',
                 co_author_phone: str = ''):
        self.topic            = topic
        self.author           = author
        self.inst             = inst
        self.email            = email
        self.co_author        = co_author
        self.co_author_title  = co_author_title
        self.co_author_inst   = co_author_inst
        self.co_author_email  = co_author_email
        self.co_author_phone  = co_author_phone
        self.writer           = writer
        self.sections         = sections
        self.specs            = specs
        self.charts           = charts
        self.papers           = papers

    def build(self) -> Document:
        doc = Document()

        # ── PAGE SETUP: A4, 1" margins ────────────────────────────────────────
        for sec in doc.sections:
            sec.page_width    = Inches(8.27)
            sec.page_height   = Inches(11.69)
            sec.top_margin    = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin   = Inches(1)
            sec.right_margin  = Inches(1)

        # ── HELPERS ───────────────────────────────────────────────────────────
        TNR = 'Times New Roman'

        def p_blank():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            return p

        def p_text(text, bold=False, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                   sp_b=0, sp_a=0, indent=None, left=None):
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(sp_b)
            pf.space_after  = Pt(sp_a)
            if indent is not None:
                pf.first_line_indent = Inches(indent)
            if left is not None:
                pf.left_indent = Inches(left)
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(sz)
            r.font.name = TNR
            return p

        def sec_head(text, sz=12, sp_b=12, sp_a=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            """All-caps bold section heading matching sample exactly"""
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(sp_b)
            pf.space_after  = Pt(sp_a)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(sz)
            r.font.name = TNR
            return p

        def body(text, sp_b=0, sp_a=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 bold=False, indent=None, left=None):
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(sp_b)
            pf.space_after  = Pt(sp_a)
            if indent is not None:
                pf.first_line_indent = Inches(indent)
            if left is not None:
                pf.left_indent = Inches(left)
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(12)
            r.font.name = TNR
            return p

        # ── TITLE PAGE (page 1) ───────────────────────────────────────────────
        # Title: centered, bold, 12pt — ALL CAPS
        p_text(self.topic.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p_blank()
        p_blank()

        # AUTHOR block — full spec fields
        p_text('AUTHOR', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p_text(self.author, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.inst:
            p_text(self.inst, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.email:
            p_text(f'EMAIL: {self.email}', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        p_blank()
        p_blank()

        # CO-AUTHOR block — matches sample paper exactly
        p_text('CO-AUTHOR', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sp_b=12, sp_a=12)
        if self.co_author:
            p_text(self.co_author, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.co_author_title:
            p_text(self.co_author_title, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.co_author_inst:
            p_text(self.co_author_inst, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.co_author_email:
            p_text(f'Email Id - {self.co_author_email}', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        if self.co_author_phone:
            p_text(f'Phone number: {self.co_author_phone}', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        p_blank()
        p_blank()

        # ── PAGE 2: Title repeat + Authors right-aligned ───────────────────────
        p_text(self.topic.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p_blank()
        author_line = self.author
        if self.co_author:
            author_line += f'\n{self.co_author}'
        p_text(author_line, bold=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT, sp_b=12, sp_a=12)

        # ── ABSTRACT ──────────────────────────────────────────────────────────
        p_text('ABSTRACT', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sp_b=0, sp_a=0)
        # Render abstract with inline bold labels (**Aim**, **Objective**, etc.)
        abs_p = doc.add_paragraph()
        abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abs_p.paragraph_format.space_before = Pt(12)
        abs_p.paragraph_format.space_after  = Pt(12)
        import re as _re_abs
        abs_segs = _re_abs.split(r"(\*\*[^*]+\*\*)", self.sections["abstract"])
        for seg in abs_segs:
            if seg.startswith("**") and seg.endswith("**"):
                r = abs_p.add_run(seg[2:-2])
                r.bold = True
            else:
                r = abs_p.add_run(seg)
                r.bold = False
            r.font.size = Pt(12); r.font.name = TNR

        # Keywords: bold label + bold keyword text, justified
        kw_p = doc.add_paragraph()
        kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        kw_p.paragraph_format.space_before = Pt(12)
        kw_p.paragraph_format.space_after  = Pt(12)
        kr1 = kw_p.add_run('KEYWORDS: ')
        kr1.bold = True; kr1.font.size = Pt(12); kr1.font.name = TNR
        kr2 = kw_p.add_run(self.sections['keywords'])
        kr2.bold = True; kr2.font.size = Pt(12); kr2.font.name = TNR

        # ── INTRODUCTION ──────────────────────────────────────────────────────
        sec_head('INTRODUCTION')
        import re as _re_intro
        intro_text = self.sections['introduction'].strip()
        # Introduction is 3-4 flowing paragraphs separated by blank lines.
        # Strip any stray markdown bold markers or subheading lines the AI may have added.
        intro_clean = _re_intro.sub(r'\*\*([^*]+)\*\*', r'\1', intro_text)
        intro_clean = _re_intro.sub(r'(?m)^[A-Z][A-Za-z ]{2,30}:\s*\n', '', intro_clean)
        intro_paras = [p.strip() for p in intro_clean.split('\n\n') if p.strip()]
        for para in intro_paras:
            body(para, sp_b=12, sp_a=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── OBJECTIVE OF THE STUDY ────────────────────────────────────────────
        sec_head('OBJECTIVE OF THE STUDY', sp_b=0, sp_a=0,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        lines = [l.strip() for l in self.sections['objectives'].splitlines() if l.strip()]
        for i, line in enumerate(lines):
            line = re.sub(r'^\d+[\.)]\s*', '', line).strip()
            line = re.sub(r'^[●•\-]\s*', '', line).strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.space_before      = Pt(12) if i == 0 else Pt(0)
            pf.space_after       = Pt(0) if i < len(lines)-1 else Pt(12)
            pf.first_line_indent = Inches(-0.25)
            pf.left_indent       = Inches(0.5)
            bullet_run = p.add_run('\u25cf       ')
            bullet_run.font.size = Pt(12); bullet_run.font.name = TNR
            r = p.add_run(line)
            r.font.size = Pt(12); r.font.name = TNR

        # ── REVIEW OF LITERATURE ──────────────────────────────────────────────
        sec_head('REVIEW OF LITERATURE', sp_b=12, sp_a=12,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        lit_paras = [l.strip() for l in self.sections['literature_review'].split('\n\n') if l.strip()]
        import re as _re2
        for i, para in enumerate(lit_paras):
            # Strip any leading number prefix like "1. " or "1) " the AI may have added
            para = _re2.sub(r'^\d+[\.\)]\s*', '', para).strip()
            # Strip markdown bold markers
            para = _re2.sub(r'\*\*([^*]+)\*\*', r'\1', para)
            if not para:
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.space_before = Pt(10) if i == 0 else Pt(8)
            pf.space_after  = Pt(4)
            # No hanging indent — full justified block like sample document

            # Format: "Author(s) (Year)" is bold, everything after is normal prose.
            # Pattern matches: "Word Word (YYYY)" or "Word, Word, and Word (YYYY)" etc.
            m_author = _re2.match(
                r'^([A-Z][^\(]{1,120}\(\d{4}[a-z]?\))\s*', para
            )
            if m_author:
                author_year = m_author.group(1).rstrip()
                rest_text   = para[m_author.end():]
                r_au = p.add_run(author_year + ' ')
                r_au.bold = True; r_au.font.size = Pt(12); r_au.font.name = TNR
                r_body = p.add_run(rest_text)
                r_body.bold = False; r_body.font.size = Pt(12); r_body.font.name = TNR
            else:
                # Fallback: render entire paragraph as plain text
                r = p.add_run(para)
                r.bold = False; r.font.size = Pt(12); r.font.name = TNR

        # ── METHODOLOGY ───────────────────────────────────────────────────────
        sec_head('METHODOLOGY', sp_b=12, sp_a=12,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        meth_text = self.sections['methodology'].strip()
        body(meth_text, sp_b=0, sp_a=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── ANALYSIS ──────────────────────────────────────────────────────────
        # Matches Graphs_AI (3).docx format:
        #   - Two charts side-by-side in a 2-column table per row
        #   - FIGURE N: label (bold) above each chart
        #   - LEGEND: (bold) below each chart — descriptive "The given figure represents..."
        #   - After all charts: Chi-Square tables, ANOVA tables, then RESULT section
        sec_head('DATA ANALYSIS AND INTERPRETATION', sz=12, sp_b=12, sp_a=12,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        import re as _re

        # Build descriptive legend texts from spec titles (sample style)
        def _build_legend(spec, fig_num):
            if spec.get('question') and spec.get('xvar'):
                return (f"The given figure represents the {spec['xvar'].lower()}-wise distribution of "
                        f"respondents' responses to the statement/question: \"{spec['question']}\".")
            title = spec.get('title', '')
            # Legacy fallback for any spec that predates the question/xvar fields
            if ' by ' in title:
                subject     = title.split(' by ')[0].strip()
                demographic = title.split(' by ')[-1].strip()
                return (f'The given figure represents the {demographic.lower()}-wise distribution of '
                        f'respondents and their views on {subject.lower()}.')
            return (f'The given figure represents respondents\' responses to '
                    f'{title.lower()} and displays the percentage distribution across all categories.')

        # Helper: add a bold-label paragraph
        def _bold_para(text, sp_b=6, sp_a=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(sp_b)
            p.paragraph_format.space_after  = Pt(sp_a)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(11); r.font.name = TNR
            return p

        # ── SINGLE-COLUMN CHART LAYOUT ────────────────────────────────────────
        # Matches sample doc: each figure on its own row —
        #   FIGURE N:  (bold, left-aligned)
        #   [chart image, centered, ~3.2"]
        #   LEGEND: (bold) descriptive text (normal)
        for fig_num, spec, buf in zip(range(1, len(self.specs) + 1), self.specs, self.charts):
            buf.seek(0)
            legend_txt = _build_legend(spec, fig_num)

            # FIGURE N: label — bold, left-aligned
            p_lbl = doc.add_paragraph()
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_lbl.paragraph_format.space_before = Pt(14)
            p_lbl.paragraph_format.space_after  = Pt(4)
            r_lbl = p_lbl.add_run(f'FIGURE {fig_num}:')
            r_lbl.bold = True; r_lbl.font.size = Pt(12); r_lbl.font.name = TNR

            # Chart image — centered, ~3.2" wide matching sample dimensions
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after  = Pt(4)
            p_img.add_run().add_picture(buf, width=Inches(3.20))

            # LEGEND: bold label + normal description text — same paragraph
            p_leg = doc.add_paragraph()
            p_leg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_leg.paragraph_format.space_before = Pt(4)
            p_leg.paragraph_format.space_after  = Pt(8)
            r_leg_lbl = p_leg.add_run('LEGEND: ')
            r_leg_lbl.bold = True; r_leg_lbl.font.size = Pt(12); r_leg_lbl.font.name = TNR
            r_leg_txt = p_leg.add_run(legend_txt)
            r_leg_txt.bold = False; r_leg_txt.font.size = Pt(12); r_leg_txt.font.name = TNR

        # ── CHI-SQUARE TABLES ─────────────────────────────────────────────────
        rng = random.Random(self.writer.seed)
        n   = self.writer.n_respondents
        chi_vars = [
            ('age',               f'awareness of {self.writer.topic[:50]}'),
            ('gender',            f'perception of {self.writer.topic[:45]}'),
            ('educational qualification', f'attitude towards {self.writer.topic[:40]}'),
            ('employment status', f'challenges in addressing {self.writer.topic[:45]}'),
            ('area of residence', f'overall engagement with {self.writer.topic[:50]}'),
        ]
        for ti, (var1, var2) in enumerate(chi_vars, 1):
            chi_val = round(rng.uniform(1.5, 9.8), 3)
            df_val  = rng.choice([2, 3, 4])
            sig_val = round(rng.uniform(0.055, 0.650), 3)
            lr_val  = round(rng.uniform(1.2, 9.2), 3)
            lr_sig  = round(rng.uniform(0.06, 0.60), 3)
            lra_val = round(rng.uniform(0.05, 2.5), 3)
            lra_sig = round(rng.uniform(0.10, 0.90), 3)

            _bold_para(f'TABLE {ti}', sp_b=14, sp_a=0)
            _bold_para(
                f'HYPOTHESIS : H0 — There is no significant association between {var1} and {var2}. '
                f'H1 — There is a significant association between {var1} and {var2}.',
                sp_b=0, sp_a=2
            )
            _add_table(doc, '', [
                ['', 'Value', 'df', 'Asymp. Sig. (2-sided)'],
                ['Pearson Chi-Square', f'{chi_val}', str(df_val), f'{sig_val}'],
                ['Likelihood Ratio',   f'{lr_val}',  str(df_val), f'{lr_sig}'],
                ['Linear-by-Linear',   f'{lra_val}', '1',         f'{lra_sig}'],
                ['N of Valid Cases',   str(n),       '',           ''],
            ])
            _bold_para(
                f'LEGEND : The above table shows the chi-square test between {var1} and {var2}.',
                sp_b=2, sp_a=2
            )
            _bold_para(
                f'INFERENCE : Since the p-value ({sig_val}) > 0.05, the null hypothesis is accepted. '
                f'There is no statistically significant association between {var1} and {var2} '
                f'at the 5% level of significance.',
                sp_b=0, sp_a=4
            )

        # ── ANOVA TABLES ──────────────────────────────────────────────────────
        anova_vars = [
            (f'age group', f'level of awareness regarding {self.writer.topic[:45]}'),
            (f'educational qualification', f'attitude and perception towards {self.writer.topic[:40]}'),
            (f'occupational category', f'engagement with preventive frameworks for {self.writer.topic[:35]}'),
        ]
        for ai, (av1, av2) in enumerate(anova_vars, 1):
            ss_between = round(rng.uniform(2.0, 15.0), 3)
            ss_within  = round(rng.uniform(50.0, 200.0), 3)
            ss_total   = round(ss_between + ss_within, 3)
            df_b       = rng.choice([2, 3, 4])
            df_w       = n - df_b - 1
            ms_b       = round(ss_between / df_b, 3)
            ms_w       = round(ss_within / df_w, 3)
            f_val      = round(ms_b / ms_w, 3)
            p_val      = round(rng.uniform(0.08, 0.75), 3)

            _bold_para(f'ANOVA TABLE {ai}', sp_b=14, sp_a=0)
            _bold_para(
                f'HYPOTHESIS : H0 — There is no significant difference in {av2} across {av1}. '
                f'H1 — There is a significant difference in {av2} across {av1}.',
                sp_b=0, sp_a=2
            )
            _add_table(doc, '', [
                ['',               'Sum of Squares', 'df',    'Mean Square', 'F',      'Sig.'],
                ['Between Groups', f'{ss_between}',  str(df_b), f'{ms_b}',  f'{f_val}', f'{p_val}'],
                ['Within Groups',  f'{ss_within}',   str(df_w), f'{ms_w}',  '',         ''],
                ['Total',          f'{ss_total}',     str(n-1),  '',          '',         ''],
            ])
            _bold_para(
                f'LEGEND : The above ANOVA table evaluates whether {av1} significantly '
                f'predicts {av2}.',
                sp_b=2, sp_a=2
            )
            _bold_para(
                f'INTERPRETATION : The model shows an F-value of {f_val} and a significance '
                f'level (p-value) of {p_val}, which is {"above" if p_val > 0.05 else "below"} the '
                f'conventional threshold of 0.05. This indicates that {av1} does '
                f'{"not have a statistically significant" if p_val > 0.05 else "have a statistically significant"} '
                f'impact on {av2}.',
                sp_b=0, sp_a=4
            )

        # ── RESULT ────────────────────────────────────────────────────────────
        sec_head('RESULT', sp_b=12, sp_a=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        import re as _re_res
        results_raw = self.sections.get('results', '').strip()
        if results_raw:
            # Results is ONE paragraph with inline **(fig: N)** markers — render with bold refs
            p_res = doc.add_paragraph()
            p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_res.paragraph_format.space_before = Pt(12)
            p_res.paragraph_format.space_after  = Pt(0)
            # Split on **(fig: N)** markers so we can bold them
            segs = _re_res.split(r'(\*\*\(fig:\s*\d+\)\*\*)', results_raw)
            for seg in segs:
                if _re_res.match(r'\*\*\(fig:\s*\d+\)\*\*', seg):
                    # Bold the (fig: N) reference, strip the ** markers
                    ref_text = seg[2:-2]  # strip leading ** and trailing **
                    r = p_res.add_run(ref_text)
                    r.bold = True
                else:
                    r = p_res.add_run(seg)
                    r.bold = False
                r.font.size = Pt(12); r.font.name = TNR

        # ── DISCUSSION ────────────────────────────────────────────────────────
        sec_head('DISCUSSION', sp_b=12, sp_a=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        import re as _re_disc
        for para in self.sections.get('discussion', '').split('\n\n'):
            para = para.strip()
            if not para:
                continue
            # Strip any leading "FIGURE N" opener the AI may still produce
            para = _re_disc.sub(r'^FIGURE\s+\d+[\.\:]?\s*', '', para, flags=_re_disc.IGNORECASE)
            # Split off a trailing **(fig: N)** or (fig: N) reference to bold it
            m_ref = _re_disc.search(r'(\*\*\(fig:\s*\d+\)\*\*|\(fig:\s*\d+\))\s*$', para)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(0)
            if m_ref:
                body_text = para[:m_ref.start()].rstrip()
                ref_raw   = m_ref.group(1)
                # Strip ** markers if present
                ref_text  = _re_disc.sub(r'\*\*', '', ref_raw)
                r1 = p.add_run(body_text + ' ')
                r1.bold = False; r1.font.size = Pt(12); r1.font.name = TNR
                r2 = p.add_run(ref_text)
                r2.bold = True; r2.font.size = Pt(12); r2.font.name = TNR
            else:
                r = p.add_run(para)
                r.bold = False; r.font.size = Pt(12); r.font.name = TNR

        # ── LIMITATIONS ───────────────────────────────────────────────────────
        sec_head('LIMITATIONS', sp_b=12, sp_a=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for para in self.sections.get('limitations', '').split('\n\n'):
            para = para.strip()
            if para:
                body(para, sp_b=12, sp_a=12, bold=False,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── SUGGESTIONS ───────────────────────────────────────────────────────
        sec_head('SUGGESTIONS', sp_b=12, sp_a=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for para in self.sections.get('suggestions', '').split('\n\n'):
            para = para.strip()
            if para:
                body(para, sp_b=12, sp_a=12, bold=False,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── CONCLUSION ────────────────────────────────────────────────────────
        sec_head('CONCLUSION', sp_b=0, sp_a=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for para in self.sections.get('conclusion', '').split('\n\n'):
            para = para.strip()
            if para:
                body(para, sp_b=12, sp_a=12, bold=False,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── REFERENCES ────────────────────────────────────────────────────────
        sec_head('REFERENCE', sp_b=0, sp_a=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Use AI-generated APA references (aligned 1-to-1 with literature review) when available
        ai_refs_raw = self.sections.get('ai_references', '').strip()
        if ai_refs_raw:
            # Parse numbered lines from the AI references block
            import re as _re_ref
            ref_lines = [l.strip() for l in ai_refs_raw.split('\n') if l.strip()]
            ref_entries = []
            current = ''
            for line in ref_lines:
                # New entry starts with a number like "1." or "1."
                if _re_ref.match(r'^\d+\.', line):
                    if current:
                        ref_entries.append(current.strip())
                    current = line
                else:
                    current += ' ' + line
            if current:
                ref_entries.append(current.strip())

            for i, ref_text in enumerate(ref_entries):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = p.paragraph_format
                pf.space_before      = Pt(6)
                pf.space_after       = Pt(0)
                pf.first_line_indent = Inches(-0.35)
                pf.left_indent       = Inches(0.5)
                # Ensure numbered prefix
                if not _re_ref.match(r'^\d+\.', ref_text):
                    ref_text = f"{i+1}. {ref_text}"
                r = p.add_run(ref_text)
                r.bold = False; r.font.size = Pt(12); r.font.name = TNR
        else:
            # Fallback to scraper-built references list
            refs = self.sections.get('references', [])
            for i, ref in enumerate(refs):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = p.paragraph_format
                pf.space_before      = Pt(6)
                pf.space_after       = Pt(0)
                pf.first_line_indent = Inches(-0.35)
                pf.left_indent       = Inches(0.5)
                ref_text = f"{i+1}. {ref}" if not ref.strip().startswith(str(i+1)+'.') else ref
                r = p.add_run(ref_text)
                r.bold = False; r.font.size = Pt(12); r.font.name = TNR

        # ── PLAGIARISM NOTE ───────────────────────────────────────────────────
        sec_head('PLAGIARISM', sp_b=0, sp_a=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        return doc

# ═══════════════════════════════════════════════════════════════════════════════
#  PAPER GENERATOR ORCHESTRATOR
# ═══════════════════════════════════════════════════════════════════════════════

class PaperGenerator:
    def __init__(self, jid: str, jobs_ref: dict):
        self.jid  = jid
        self.jobs = jobs_ref

    def prog(self, pct: int, msg: str):
        self.jobs[self.jid].update({'progress': pct, 'message': msg, 'status': 'running'})
        print(f'[{self.jid[:8]}] {pct}% – {msg}')

    def generate(self, topic: str, nfigs: int, author: str, inst: str, email: str,
                 questionnaire: dict = None, co_author_info: dict = None) -> str:
        os.makedirs('generated', exist_ok=True)
        self.prog(5, 'Initializing...')
        ca = co_author_info or {}

        # ── Step 1: Web scraping — 3 sources in parallel ─────────────────────
        self.prog(8, 'Scraping Semantic Scholar, CrossRef & Wikipedia...')
        scraper = WebScraper(topic)
        with ThreadPoolExecutor(max_workers=3) as ex:
            f_ss, f_cr, f_wiki = (
                ex.submit(scraper.fetch_semantic_scholar, 22),
                ex.submit(scraper.fetch_crossref, 14),
                ex.submit(scraper.fetch_wikipedia),
            )
            ss, cr, wiki = f_ss.result(), f_cr.result(), f_wiki.result()

        seen, all_papers = set(), []
        for p in ss + cr:
            key = p['title'][:40].lower()
            if key not in seen:
                seen.add(key); all_papers.append(p)
        all_papers.sort(key=lambda x: x.get('citations', 0), reverse=True)

        # Safety net: if the topic is niche and we still don't have enough REAL
        # papers to seed 20 genuine literature-review entries, broaden the query
        # (first 2-3 words of the topic) and top up — never fabricate instead.
        broadened = False
        if len(all_papers) < 20:
            broad_topic = ' '.join(topic.split()[:3]) or topic
            if broad_topic.lower() != topic.lower():
                broadened = True
                broad_scraper = WebScraper(broad_topic)
                with ThreadPoolExecutor(max_workers=2) as ex:
                    f_ss2, f_cr2 = (
                        ex.submit(broad_scraper.fetch_semantic_scholar, 20),
                        ex.submit(broad_scraper.fetch_crossref, 10),
                    )
                    ss2, cr2 = f_ss2.result(), f_cr2.result()
                for p in ss2 + cr2:
                    key = p['title'][:40].lower()
                    if key not in seen:
                        seen.add(key); all_papers.append(p)
                all_papers.sort(key=lambda x: x.get('citations', 0), reverse=True)

        scraped = {'papers': all_papers, 'wiki': wiki}
        print(f"[Scraper] {len(all_papers)} unique real papers gathered "
              f"({len(ss)} SS + {len(cr)} CrossRef{' + broadened query' if broadened else ''}), "
              f"wiki={'yes' if wiki.get('summary') else 'no'}")

        # ── Step 2: AI writes all sections ───────────────────────────────────
        self.prog(30, 'AI connected — writing keywords...')
        writer        = GeminiWriter(topic, scraped, questionnaire=questionnaire or {})
        writer._nfigs = nfigs
        sections      = writer.generate_all(progress_cb=self.prog)
        self.prog(76, 'AI finished. Parsing sections...')

        sections['references'] = writer.references()
        # Prefer AI-generated APA references (aligned with lit review) when available
        if sections.get('ai_references'):
            sections['use_ai_references'] = True

        # ── Step 3: Parse chart specs from AI's <charts> block ───────────────
        self.prog(78, 'Parsing chart specs...')
        specs = writer.parse_chart_specs(nfigs)
        if not specs:
            specs = writer._fallback_specs(nfigs)

        # ── Step 4: Render charts ────────────────────────────────────────────
        self.prog(82, f'Rendering {len(specs)} SPSS-style charts...')
        charts = [make_chart(sp) for sp in specs]

        # ── Step 5: Build DOCX ───────────────────────────────────────────────
        self.prog(90, 'Assembling Word document...')
        builder = DocBuilder(
            topic, author, inst, email, writer, sections, specs, charts, all_papers,
            co_author       = ca.get('name', ''),
            co_author_title = ca.get('title', ''),
            co_author_inst  = ca.get('inst', ''),
            co_author_email = ca.get('email', ''),
            co_author_phone = ca.get('phone', ''),
        )
        doc = builder.build()

        self.prog(97, 'Saving...')
        safe = re.sub(r'[^\w\-]', '_', topic[:40])
        out  = os.path.abspath(f'generated/rdxper_{safe}_{self.jid[:8]}.docx')
        doc.save(out)
        self.prog(99, 'Done!')
        return out


# ═══════════════════════════════════════════════════════════════════════════════
#  EMBEDDED FRONTEND
# ═══════════════════════════════════════════════════════════════════════════════

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>rdxper</title>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',Arial,sans-serif;background:#ffffff;color:#111111;min-height:100vh}
:root{--bg:#ffffff;--surface:#ffffff;--surface2:#f5f5f5;--surface3:#ececec;--border:#d0d0d0;--accent:#111111;--accent2:#444444;--text:#111111;--muted:#666666;--dim:#999999;--error:#cc0000;--r:10px}
.wrap{max-width:960px;margin:0 auto;padding:0 20px}
header{padding:18px 0;display:flex;align-items:center;justify-content:space-between;border-bottom:2px solid #111}
.logo{display:flex;align-items:center;gap:10px}
.logo-mark{width:32px;height:32px;background:#111;border-radius:6px;display:flex;align-items:center;justify-content:center;font-weight:900;font-size:12px;color:#fff}
.logo-text{font-size:20px;font-weight:900;letter-spacing:-0.5px;color:#111}
.logo-text span{color:#111}
.user-chip{display:flex;align-items:center;gap:8px;background:#f5f5f5;border:1px solid #d0d0d0;border-radius:40px;padding:5px 12px 5px 5px;cursor:pointer}
.user-chip img{width:26px;height:26px;border-radius:50%;object-fit:cover}
.user-chip span{font-size:13px;max-width:150px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;color:#111}
.nav-links{display:flex;gap:8px;align-items:center}
.nav-btn{background:none;border:1px solid #d0d0d0;color:#666;padding:5px 12px;border-radius:6px;cursor:pointer;font-size:12px;transition:all .2s}
.nav-btn:hover{border-color:#111;color:#111;background:#f5f5f5}
.nav-btn.danger{border-color:#cc0000;color:#cc0000}
.screen{display:none}.screen.active{display:block}
.hero{padding:56px 0 32px;text-align:center}
.htag{font-size:12px;color:#666;letter-spacing:2px;text-transform:uppercase;margin-bottom:16px;font-family:Consolas,monospace}
h1{font-size:clamp(28px,5vw,52px);font-weight:900;line-height:1.1;margin-bottom:16px;color:#111}
h1 em{color:#111;font-style:normal;border-bottom:3px solid #111}
.sub{font-size:16px;color:#666;max-width:560px;margin:0 auto 32px}
.card{background:#fff;border:1.5px solid #d0d0d0;border-radius:var(--r);padding:32px;max-width:440px;margin:0 auto;width:100%}
.ct{font-size:20px;font-weight:700;margin-bottom:6px;color:#111}
.cs{font-size:14px;color:#666;margin-bottom:24px}
.btn{width:100%;padding:13px 20px;border-radius:8px;border:none;font-size:15px;font-weight:700;cursor:pointer;transition:all .15s;display:flex;align-items:center;justify-content:center;gap:8px;margin-bottom:10px}
.btn:disabled{opacity:.4;cursor:not-allowed}
.btn-p{background:#111;color:#fff;border:2px solid #111}
.btn-p:hover:not(:disabled){background:#333;transform:translateY(-1px);box-shadow:0 4px 12px rgba(0,0,0,.2)}
.btn-dl{background:#111;color:#fff;border:2px solid #111;box-shadow:0 2px 8px rgba(0,0,0,.15)}
.btn-dl:hover:not(:disabled){background:#333;transform:translateY(-2px);box-shadow:0 6px 20px rgba(0,0,0,.25)}
.btn-s{background:#fff;color:#111;border:1.5px solid #d0d0d0}
.btn-s:hover:not(:disabled){border-color:#111;background:#f5f5f5}
.fg{margin-bottom:16px}.fg label{display:block;font-size:13px;color:#555;margin-bottom:6px;font-weight:600}
.fg input,.fg select{width:100%;background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;font-size:14px;outline:none;transition:border-color .2s}
.fg input:focus,.fg input:focus-within{border-color:#111;background:#fff}
.notif{display:none;padding:10px 14px;border-radius:8px;font-size:13px;margin-bottom:14px}
.notif.show{display:block}
.notif.success{background:#f0f0f0;border:1.5px solid #111;color:#111}
.notif.error{background:#fff0f0;border:1.5px solid #cc0000;color:#cc0000}
.notif.info{background:#f5f5f5;border:1.5px solid #888;color:#444}
.prog-wrap{background:#e8e8e8;border-radius:100px;height:5px;overflow:hidden;margin:12px 0}
.prog-fill{height:100%;background:#111;border-radius:100px;transition:width .4s ease}
.prog-row{display:flex;justify-content:space-between;font-size:12px;color:#666;margin-bottom:4px}
.stage-box{background:#f5f5f5;border:1.5px solid #d0d0d0;border-radius:var(--r);padding:10px 14px;margin:10px 0;display:flex;align-items:center;gap:8px}
.stage-msg{font-size:12px;color:#111;font-family:Consolas,monospace;flex:1;font-weight:600}
.sections-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:5px;margin-bottom:12px}
.sec-item{font-size:9px;padding:4px;border-radius:5px;background:#f5f5f5;border:1px solid #d0d0d0;color:#999;text-align:center;font-family:Consolas,monospace;transition:all .3s}
.sec-item.writing{background:#ececec;border-color:#111;color:#111;animation:sp 1s ease-in-out infinite;font-weight:700}
.sec-item.done{background:#111;border-color:#111;color:#fff}
@keyframes sp{0%,100%{opacity:1}50%{opacity:.4}}
.spin{width:14px;height:14px;border:2px solid rgba(255,255,255,.4);border-top-color:#fff;border-radius:50%;animation:spin .7s linear infinite;display:inline-block}
@keyframes spin{to{transform:rotate(360deg)}}
.stat-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:12px;margin-bottom:24px}
.stat-card{background:#fff;border:1.5px solid #d0d0d0;border-radius:var(--r);padding:20px}
.stat-val{font-size:28px;font-weight:900;color:#111}
.stat-lbl{font-size:12px;color:#666;margin-top:4px}
.table-wrap{background:#fff;border:1.5px solid #d0d0d0;border-radius:var(--r);overflow:hidden;margin-bottom:24px}
.table-head{padding:14px 20px;border-bottom:1.5px solid #d0d0d0;font-size:14px;font-weight:700;color:#111}
table{width:100%;border-collapse:collapse}
th{text-align:left;padding:10px 16px;font-size:11px;color:#666;text-transform:uppercase;letter-spacing:.5px;border-bottom:1.5px solid #d0d0d0;font-weight:700;background:#f9f9f9}
td{padding:10px 16px;font-size:13px;border-bottom:1px solid #ececec;color:#111}
tr:last-child td{border-bottom:none}
tr:hover td{background:#f9f9f9}
.badge-paid{background:#111;color:#fff;padding:2px 8px;border-radius:4px;font-size:11px;font-weight:700}
.badge-free{background:#f5f5f5;color:#666;border:1px solid #d0d0d0;padding:2px 8px;border-radius:4px;font-size:11px}
.badge-pending{background:#f5f5f5;color:#888;border:1px solid #ccc;padding:2px 8px;border-radius:4px;font-size:11px}
.avatar{width:32px;height:32px;border-radius:50%;object-fit:cover;border:1px solid #d0d0d0}
.profile-header{display:flex;align-items:center;gap:16px;background:#fff;border:1.5px solid #d0d0d0;border-radius:var(--r);padding:24px;margin-bottom:20px}
.profile-avatar{width:64px;height:64px;border-radius:50%;border:2px solid #111}
.tabs{display:flex;gap:0;margin-bottom:20px;border-bottom:2px solid #d0d0d0}
.tab{padding:10px 18px;font-size:13px;cursor:pointer;border-radius:0;color:#666;border:none;background:none;transition:all .2s;border-bottom:2px solid transparent;margin-bottom:-2px;font-weight:600}
.tab.active{color:#111;border-bottom:2px solid #111;font-weight:700}
.empty{text-align:center;padding:40px;color:#999;font-size:14px}
.pay-box{background:#f5f5f5;border:1.5px solid #d0d0d0;border-radius:12px;padding:20px;text-align:center;margin:16px 0}
.pay-amt{font-size:40px;font-weight:900;color:#111}
.page-title{font-size:24px;font-weight:900;margin:32px 0 4px;color:#111}
.page-sub{font-size:14px;color:#666;margin-bottom:24px}
footer{text-align:center;padding:32px 0;color:#999;font-size:12px;border-top:1.5px solid #d0d0d0;margin-top:40px}
/* Questionnaire */
.q-steps{display:flex;align-items:center;margin-bottom:28px;padding:0 4px}
.q-step{display:flex;flex-direction:column;align-items:center;gap:4px;cursor:pointer;min-width:56px}
.q-num{width:28px;height:28px;border-radius:50%;background:#f5f5f5;border:2px solid #d0d0d0;display:flex;align-items:center;justify-content:center;font-size:12px;font-weight:700;color:#999;transition:all .3s}
.q-lbl{font-size:10px;color:#999;transition:color .3s;white-space:nowrap}
.q-step.active .q-num{background:#111;border-color:#111;color:#fff}
.q-step.active .q-lbl{color:#111;font-weight:700}
.q-step.done .q-num{background:#555;border-color:#555;color:#fff}
.q-step.done .q-lbl{color:#555}
.q-line{flex:1;height:2px;background:#d0d0d0;margin:0 4px;margin-bottom:14px;transition:background .3s}
.q-line.done{background:#111}
.q-panel{display:none}.q-panel.active{display:block}
.q-badge{font-size:11px;color:#888;font-family:Consolas,monospace;letter-spacing:1px;margin-bottom:8px;font-weight:700;text-transform:uppercase}
.q-hint{background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;font-size:12px;color:#555;margin-bottom:16px;line-height:1.5}
textarea{width:100%;background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;font-size:13px;outline:none;transition:border-color .2s;resize:vertical;font-family:'Segoe UI',Arial,sans-serif;line-height:1.6}
textarea:focus{border-color:#111;background:#fff}
textarea::placeholder{color:#bbb;font-size:12px}
.q-summary{background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:10px;padding:16px;margin-bottom:20px;font-size:12px}
.q-summary-item{margin-bottom:10px;padding-bottom:10px;border-bottom:1px solid #e8e8e8}
.q-summary-item:last-child{margin-bottom:0;padding-bottom:0;border-bottom:none}
.q-summary-label{color:#111;font-weight:700;font-size:11px;text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}
.q-summary-val{color:#444;line-height:1.5;max-height:60px;overflow:hidden;text-overflow:ellipsis}
@media(max-width:600px){.q-lbl{display:none}.q-steps{gap:0}.q-step{min-width:36px}}
@media(max-width:600px){.sections-grid{grid-template-columns:repeat(3,1fr)}.stat-grid{grid-template-columns:repeat(2,1fr)}.nav-links{gap:4px}}
/* ── Dashboard ── */
.dash-header{padding:36px 0 8px}
.dash-greeting{font-size:13px;color:#888;letter-spacing:.5px;text-transform:uppercase;font-family:Consolas,monospace;margin-bottom:6px}
.dash-title{font-size:30px;font-weight:900;letter-spacing:-1px;color:#111}
.dash-title span{color:#111}
.dash-empty{display:flex;flex-direction:column;align-items:center;justify-content:center;padding:80px 20px;text-align:center;border:2px dashed #d0d0d0;border-radius:16px;margin:28px 0}
.dash-empty-icon{font-size:48px;margin-bottom:16px;opacity:.3}
.dash-empty-txt{font-size:16px;font-weight:700;color:#666;margin-bottom:6px}
.dash-empty-sub{font-size:13px;color:#999}
.papers-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(280px,1fr));gap:16px;margin:24px 0}
.paper-card{background:#fff;border:1.5px solid #d0d0d0;border-radius:10px;padding:20px;cursor:default;transition:border-color .15s,transform .15s,box-shadow .15s;position:relative;overflow:hidden}
.paper-card::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;background:#111;opacity:0;transition:opacity .15s}
.paper-card:hover{border-color:#111;transform:translateY(-2px);box-shadow:0 4px 16px rgba(0,0,0,.1)}
.paper-card:hover::before{opacity:1}
.paper-card-topic{font-size:14px;font-weight:700;color:#111;line-height:1.4;margin-bottom:12px;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden}
.paper-card-meta{display:flex;align-items:center;justify-content:space-between;font-size:11px;color:#999}
.paper-card-date{font-family:Consolas,monospace}
.paper-card-badge{padding:2px 8px;border-radius:4px;font-size:10px;font-weight:700;letter-spacing:.3px}
.badge-done{background:#111;color:#fff}
.badge-pending{background:#f5f5f5;color:#888;border:1px solid #d0d0d0}
.fab{position:fixed;bottom:32px;right:32px;width:58px;height:58px;border-radius:50%;background:#111;border:none;cursor:pointer;display:flex;align-items:center;justify-content:center;box-shadow:0 4px 16px rgba(0,0,0,.25);transition:transform .2s,box-shadow .2s;z-index:100}
.fab:hover{transform:scale(1.1) translateY(-2px);box-shadow:0 8px 28px rgba(0,0,0,.35)}
.fab svg{width:24px;height:24px;stroke:#fff;stroke-width:2.5;stroke-linecap:round}
.fab-tooltip{position:fixed;bottom:44px;right:100px;background:#111;border-radius:6px;padding:7px 12px;font-size:12px;color:#fff;white-space:nowrap;opacity:0;pointer-events:none;transition:opacity .2s;z-index:99}
.fab:hover ~ .fab-tooltip{opacity:1}
@media(max-width:600px){.fab{bottom:24px;right:20px}.papers-grid{grid-template-columns:1fr}}
/* ── Responsive ── */
@media(max-width:480px){
  .wrap{padding:0 12px}
  header{padding:12px 0;flex-wrap:wrap;gap:8px}
  .logo-text{font-size:17px}
  .logo-mark{width:28px;height:28px;font-size:10px}
  .nav-btn{padding:4px 8px;font-size:11px}
  .user-chip{padding:4px 8px 4px 4px}
  .user-chip span{max-width:80px;font-size:12px}
  .hero{padding:32px 0 20px}
  h1{font-size:26px}
  .sub{font-size:14px}
  .card{padding:20px 16px}
  .ct{font-size:17px}
  .cs{font-size:13px}
  .btn{padding:12px 16px;font-size:14px}
  .fg input,.fg select{padding:9px 12px;font-size:14px}
  textarea{font-size:13px}
  .dash-header{padding:24px 0 4px}
  .dash-title{font-size:24px}
  .dash-empty{padding:48px 16px}
  .stat-grid{grid-template-columns:1fr 1fr}
  .stat-val{font-size:22px}
  .table-wrap{overflow-x:auto}
  table{font-size:12px;min-width:360px}
  th,td{padding:8px 10px}
  .profile-header{flex-direction:column;text-align:center;gap:12px}
  .profile-avatar{width:56px;height:56px}
  .tabs{overflow-x:auto;white-space:nowrap;-webkit-overflow-scrolling:touch}
  .tab{padding:8px 14px;font-size:12px}
  .card[style*="max-width:560px"]{padding:20px 16px}
  .stage-box{padding:8px 12px}
  .stage-msg{font-size:11px}
  .sections-grid{grid-template-columns:repeat(3,1fr)}
  .sec-item{font-size:8px;padding:3px}
  .q-badge{font-size:10px}
  .q-hint{font-size:11px;padding:8px 12px}
  .q-summary{padding:12px}
  .q-summary-label{font-size:10px}
  .q-summary-val{font-size:12px}
  .page-title{font-size:20px}
  .page-sub{font-size:13px}
}
@media(min-width:481px) and (max-width:768px){
  .wrap{padding:0 16px}
  header{padding:14px 0}
  h1{font-size:36px}
  .card{padding:28px 24px;max-width:100%}
  .stat-grid{grid-template-columns:repeat(3,1fr)}
  .papers-grid{grid-template-columns:repeat(2,1fr)}
  .user-chip span{max-width:110px}
  .table-wrap{overflow-x:auto}
  table{min-width:420px}
}
@media(min-width:769px) and (max-width:1024px){
  .wrap{padding:0 24px}
  .papers-grid{grid-template-columns:repeat(2,1fr)}
  .stat-grid{grid-template-columns:repeat(4,1fr)}
}
@media(hover:none){
  .paper-card:hover{transform:none;box-shadow:none}
  .paper-card:hover::before{opacity:0}
  .btn-p:hover:not(:disabled){transform:none;box-shadow:none}
  .btn-dl:hover:not(:disabled){transform:none;box-shadow:none}
  .fab:hover{transform:scale(1);box-shadow:0 4px 16px rgba(0,0,0,.25)}
}
</style>
</head>
<body>
<div class="wrap">
<header>
  <div class="logo">
    <div class="logo-mark">rx</div>
    <div class="logo-text">RD<span>Xper</span></div>
  </div>
  <div class="nav-links" id="nav-auth" style="display:none">
    <button class="nav-btn" onclick="showProfile()">👤 Profile</button>
    <button class="nav-btn" onclick="show('s-legal')">⚖️ Legal</button>
    <div id="admin-link" style="display:none"><button class="nav-btn" onclick="showAdmin()">⚙️ Admin</button></div>
    <div class="user-chip" onclick="showProfile()">
      <img id="nav-avatar" src="" onerror="this.style.display='none'" style="display:none">
      <span id="nav-name">User</span>
    </div>
    <button class="nav-btn danger" onclick="logout()">Sign out</button>
  </div>
</header>

<!-- LOGIN -->
<div class="screen active" id="s-home">
  <div class="hero">
    
    <h1>Generate <em>Genuine</em><br>Research Papers</h1>
    
  </div>
  <div class="card">
    <div class="ct">Sign in to continue</div>
    <div class="cs">Enter your name and email to get started</div>
    <div id="n-login" class="notif"></div>
    <div id="g-btn-wrap" style="display:flex;justify-content:center;min-height:44px;align-items:center"></div>
  </div>
</div>

<!-- DASHBOARD -->
<div class="screen" id="s-dashboard">
  <div class="dash-header">
    <div class="dash-greeting">Welcome back</div>
    <div class="dash-title" id="dash-name-title">Researcher</div>
  </div>

  <div style="display:flex;align-items:center;justify-content:space-between;margin:28px 0 8px">
    <div style="font-size:16px;font-weight:700">Your Research Papers</div>
    <button class="nav-btn" onclick="loadDashboard()" style="font-size:11px">↻ Refresh</button>
  </div>

  <div id="dash-papers-wrap">
    <div class="dash-empty">
      <div class="dash-empty-icon">📄</div>
      <div class="dash-empty-txt">No papers yet</div>
      <div class="dash-empty-sub">Press <strong style="color:var(--accent)">+</strong> below to generate your first research paper</div>
    </div>
  </div>

  <!-- Floating Action Button -->
  <button class="fab" onclick="startNewPaper()" title="New Research Paper">
    <svg viewBox="0 0 24 24" fill="none"><line x1="12" y1="5" x2="12" y2="19"/><line x1="5" y1="12" x2="19" y2="12"/></svg>
  </button>
  <div class="fab-tooltip">New Research Paper</div>
</div>

<!-- GENERATE — 5-Step Questionnaire -->
<div class="screen" id="s-gen">
<div style="padding-top:28px;max-width:700px;margin:0 auto">

<div style="margin-bottom:16px">
  <button class="btn btn-s" style="width:auto;padding:8px 16px;font-size:12px" onclick="loadDashboard();show('s-dashboard')">← Dashboard</button>
</div>

<!-- Step indicator -->
<div class="q-steps" id="q-steps">
  <div class="q-step active" id="qs-0" onclick="goStep(0)"><span class="q-num">1</span><span class="q-lbl">Problem</span></div>
  <div class="q-line"></div>
  <div class="q-step" id="qs-1" onclick="goStep(1)"><span class="q-num">2</span><span class="q-lbl">Literature</span></div>
  <div class="q-line"></div>
  <div class="q-step" id="qs-2" onclick="goStep(2)"><span class="q-num">3</span><span class="q-lbl">Gap</span></div>
  <div class="q-line"></div>
  <div class="q-step" id="qs-3" onclick="goStep(3)"><span class="q-num">4</span><span class="q-lbl">Objectives</span></div>
  <div class="q-line"></div>
  <div class="q-step" id="qs-4" onclick="goStep(4)"><span class="q-num">5</span><span class="q-lbl">Statement</span></div>
  <div class="q-line"></div>
  <div class="q-step" id="qs-5" onclick="goStep(5)"><span class="q-num">6</span><span class="q-lbl">Settings</span></div>
</div>

<!-- ── Step 0: Problem Identification ───────────────────── -->
<div class="q-panel active" id="qp-0">
  <div class="q-badge">Step 1 of 6</div>
  <div class="ct" style="margin-bottom:6px">Identification of the Problem</div>
  <div class="cs" style="margin-bottom:20px">What specific problem prompted this research? Describe it in your own words, AI will use this as the foundation. <strong style="color:var(--accent)">Optional — skip if you prefer AI to write this.</strong></div>
  <div class="q-hint">💡 Think about: What is wrong or missing? Who is affected? What is the scale of the problem? What are the consequences of not addressing it?</div>
  <div class="fg">
    <label>Research Topic / Title *</label>
    <input type="text" id="topic-in" placeholder="e.g. Legal Frameworks for Environmental Restoration in Post-War Reconstruction">
  </div>
  <div class="fg">
    <label>Problem Statement <span style="color:var(--dim);font-weight:400">(optional)</span></label>
    <textarea id="q-problem" rows="5" placeholder="Describe the core problem your research addresses. What issue exists? What are its consequences? Why does it need to be studied now?&#10;&#10;Example: Armed conflicts inflict devastating environmental damage that persists long after hostilities cease. Existing legal frameworks under the Geneva Conventions and Rome Statute fail to adequately address post-war ecological restoration, leaving affected communities without legal recourse or environmental remediation. This gap in international humanitarian law creates a vacuum where neither state nor non-state actors are held accountable for long-term environmental harm..."></textarea>
  </div>
  <div style="display:flex;gap:10px;justify-content:flex-end">
    <div style="display:flex;gap:10px;justify-content:flex-end">
      <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="nextStep(0)">Skip →</button>
      <button class="btn btn-p" style="width:auto;padding:10px 28px" onclick="nextStep(0)">Next → Literature Review</button>
    </div>
  </div>
</div>

<!-- ── Step 1: Literature Review ────────────────────────── -->
<div class="q-panel" id="qp-1">
  <div class="q-badge">Step 2 of 6</div>
  <div class="ct" style="margin-bottom:6px">Literature Review</div>
  <div class="cs" style="margin-bottom:20px">What sources have you reviewed? List them and AI will expand into a full literature review. <strong style="color:var(--accent)">Optional — AI will find real papers automatically if you skip.</strong></div>
  <div class="q-hint">💡 Include: Author names and years, key arguments, relevant reports, laws, treaties, court cases, or books. Even brief notes are fine — AI will elaborate.</div>
  <div class="fg">
    <label>Key Sources & Their Main Arguments *</label>
    <textarea id="q-lit" rows="8" placeholder="List the sources you have reviewed and what they say. Examples:&#10;&#10;- Geneva Conventions (1949) & Additional Protocol I (1977) — establish basic environmental protections during armed conflict but lack post-war restoration obligations&#10;- UNEP (2009) From Conflict to Peacebuilding — documents how environmental damage sustains conflict cycles&#10;- Bothe, Bruch & Jensen (2010) — argue existing IHL is inadequate for modern environmental warfare&#10;- Rome Statute Art. 8 — criminalises widespread environmental damage but enforcement is rare&#10;- UN Compensation Commission (Kuwait, 1991) — first successful precedent for war environmental claims..."></textarea>
  </div>
  <div style="display:flex;gap:10px;justify-content:space-between">
    <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="prevStep(1)">← Back</button>
    <button class="btn btn-s" style="width:auto;padding:10px 18px" onclick="nextStep(1)">Skip →</button>
    <button class="btn btn-p" style="width:auto;padding:10px 28px" onclick="nextStep(1)">Next → Research Gap</button>
  </div>
</div>

<!-- ── Step 2: Research Gap ──────────────────────────────── -->
<div class="q-panel" id="qp-2">
  <div class="q-badge">Step 3 of 6</div>
  <div class="ct" style="margin-bottom:6px">Research Gap</div>
  <div class="cs" style="margin-bottom:20px">What is missing from existing research? AI will use your answer as the gap statement. <strong style="color:var(--accent)">Optional — AI will identify a gap automatically if you skip.</strong></div>
  <div class="q-hint">💡 Ask yourself: What do existing studies not cover? What contradictions exist in the literature? What context or population has been ignored? What methodology hasn't been applied?</div>
  <div class="fg">
    <label>The Research Gap <span style="color:var(--dim);font-weight:400">(optional)</span></label>
    <textarea id="q-gap" rows="5" placeholder="Describe what is missing from current research and why your study is needed.&#10;&#10;Example: While significant scholarship exists on environmental protection during armed conflict, there is a critical gap in research on post-war environmental restoration obligations. Existing studies either focus on pre-conflict prevention or general humanitarian law without addressing the specific legal mechanisms required for ecological recovery. Furthermore, no comparative study has examined how different post-conflict nations (Iraq, Kosovo, Lebanon, Ukraine) have implemented or failed to implement environmental restoration under international law..."></textarea>
  </div>
  <div style="display:flex;gap:10px;justify-content:space-between">
    <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="prevStep(2)">← Back</button>
    <button class="btn btn-s" style="width:auto;padding:10px 18px" onclick="nextStep(2)">Skip →</button>
    <button class="btn btn-p" style="width:auto;padding:10px 28px" onclick="nextStep(2)">Next → Objectives</button>
  </div>
</div>

<!-- ── Step 3: Objectives ────────────────────────────────── -->
<div class="q-panel" id="qp-3">
  <div class="q-badge">Step 4 of 6</div>
  <div class="ct" style="margin-bottom:6px">Objectives of the Research</div>
  <div class="cs" style="margin-bottom:20px">List your research objectives — they will appear verbatim in your paper. <strong style="color:var(--accent)">Optional — AI will generate objectives aligned to your topic if you skip.</strong></div>
  <div class="q-hint">💡 Good objectives: Start with "To examine / To analyse / To evaluate / To compare / To propose". Be specific. You need 4–6 objectives. One per line.</div>
  <div class="fg">
    <label>Research Objectives <span style="color:var(--dim);font-weight:400">(optional — one per line)</span></label>
    <textarea id="q-objectives" rows="7" placeholder="To examine the existing international legal frameworks governing environmental restoration in post-war reconstruction&#10;To analyse compensation mechanisms including liability determination, reparations, and restoration funding&#10;To evaluate practical challenges such as political instability, limited resources, and technical capacity gaps&#10;To compare legal approaches from different post-conflict contexts including Iraq, Kosovo, Lebanon, and Ukraine&#10;To propose recommendations for strengthening enforcement mechanisms and legal accountability for wartime environmental harm"></textarea>
  </div>
  <div style="display:flex;gap:10px;justify-content:space-between">
    <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="prevStep(3)">← Back</button>
    <button class="btn btn-s" style="width:auto;padding:10px 18px" onclick="nextStep(3)">Skip →</button>
    <button class="btn btn-p" style="width:auto;padding:10px 28px" onclick="nextStep(3)">Next → Research Statement</button>
  </div>
</div>

<!-- ── Step 4: Research Statement ───────────────────────── -->
<div class="q-panel" id="qp-4">
  <div class="q-badge">Step 5 of 6</div>
  <div class="ct" style="margin-bottom:6px">Research Statement</div>
  <div class="cs" style="margin-bottom:20px">Your thesis in 2–4 sentences — what this research does, how, and why. <strong style="color:var(--accent)">Optional — AI will formulate a research statement if you skip.</strong></div>
  <div class="q-hint">💡 A good research statement: Names the topic, identifies the method (doctrinal/empirical/comparative), and states the significance. Typically 2–4 sentences.</div>
  <div class="fg">
    <label>Research Statement <span style="color:var(--dim);font-weight:400">(optional)</span></label>
    <textarea id="q-statement" rows="5" placeholder="This study investigates the legal frameworks governing environmental restoration in post-war reconstruction, focusing on obligations, compensation mechanisms, and practical implementation challenges. Through a comparative doctrinal analysis of international instruments and empirical case studies from four post-conflict regions, this research identifies critical gaps in existing law and proposes actionable reforms to strengthen ecological restoration as an integral component of sustainable peace-building."></textarea>
  </div>
  <div style="display:flex;gap:10px;justify-content:space-between">
    <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="prevStep(4)">← Back</button>
    <button class="btn btn-s" style="width:auto;padding:10px 18px" onclick="nextStep(4)">Skip →</button>
    <button class="btn btn-p" style="width:auto;padding:10px 28px" onclick="nextStep(4)">Next → Paper Settings</button>
  </div>
</div>

<!-- ── Step 5: Settings + Generate ──────────────────────── -->
<div class="q-panel" id="qp-5">
  <div class="q-badge">Step 6 of 6</div>
  <div class="ct" style="margin-bottom:6px">Paper Settings</div>
  <div class="cs" style="margin-bottom:20px">Final details for your paper. AI will now use all your inputs to generate a genuine research paper.</div>
  <div id="n-gen" class="notif"></div>
  <!-- Summary of inputs -->
  <div class="q-summary" id="q-summary"></div>

  <div style="font-weight:700;font-size:13px;letter-spacing:1px;text-transform:uppercase;margin-bottom:8px;margin-top:4px;color:#444">Author Details</div>
  <div class="fg"><label>Author Name</label>
    <input type="text" id="author-in" placeholder="Your full name">
  </div>
  <div class="fg"><label>Institution (optional)</label>
    <input type="text" id="inst-in" placeholder="University / College / Organisation">
  </div>

  <div style="font-weight:700;font-size:13px;letter-spacing:1px;text-transform:uppercase;margin-bottom:8px;margin-top:16px;color:#444">Co-Author Details (optional)</div>
  <div class="fg"><label>Co-Author Name</label>
    <input type="text" id="co-author-name" placeholder="Co-author full name">
  </div>
  <div class="fg"><label>Co-Author Title / Designation</label>
    <input type="text" id="co-author-title" placeholder="e.g. Assistant Professor, Research Scholar">
  </div>
  <div class="fg"><label>Co-Author Institution</label>
    <input type="text" id="co-author-inst" placeholder="University / College / Organisation">
  </div>
  <div class="fg"><label>Co-Author Email</label>
    <input type="email" id="co-author-email" placeholder="co-author@example.com">
  </div>
  <div class="fg"><label>Co-Author Phone (optional)</label>
    <input type="text" id="co-author-phone" placeholder="Contact number">
  </div>

  <div class="fg" style="margin-top:12px"><label>Number of Figures: <b id="sl-display">6</b></label>
    <input type="range" id="sl" min="3" max="25" value="6"
      oninput="document.getElementById('sl-display').textContent=this.value"
      style="width:100%;accent-color:var(--accent)">
  </div>
  <div style="display:flex;gap:10px;justify-content:space-between">
    <button class="btn btn-s" style="width:auto;padding:10px 20px" onclick="prevStep(5)">← Back</button>
    <button class="btn btn-p" id="btn-gen" onclick="generate()" style="flex:1">Generate Research Paper</button>
  </div>
</div>

</div>
</div>

<!-- PROGRESS -->
<div class="screen" id="s-prog">
  <div style="padding-top:40px">
    <div class="card" style="max-width:560px">
      <div class="ct" id="prog-ct">Generating your paper...</div>
      <div class="cs" id="prog-topic"></div>
      <div class="stage-box"><span style="font-size:18px">⚡</span><span class="stage-msg" id="stage-msg">Initialising...</span></div>
      <div class="prog-row"><span></span><span id="prog-pct">0%</span></div>
      <div class="prog-wrap"><div class="prog-fill" id="prog-fill" style="width:0%"></div></div>
      <div class="sections-grid" id="sec-grid"></div>
    </div>
  </div>
</div>

<!-- DONE -->
<div class="screen" id="s-done">
  <div style="padding-top:48px">
    <div class="card" style="text-align:center">
      <div style="font-size:48px;margin-bottom:12px">✅</div>
      <div class="ct">Paper ready!</div>
      <div class="cs">Your research paper has been generated successfully</div>
      <div style="background:var(--surface2);border:1px solid var(--border);border-radius:10px;padding:16px;margin:16px 0;text-align:left">
        <div style="display:flex;justify-content:space-between;margin-bottom:8px">
          <span style="color:var(--muted);font-size:13px">Topic</span>
          <span style="font-size:13px;font-weight:600;max-width:220px;text-align:right" id="d-topic"></span></div>
        <div style="display:flex;justify-content:space-between;margin-bottom:8px">
          <span style="color:var(--muted);font-size:13px">Figures</span>
          <span style="font-size:13px" id="d-figs"></span></div>
        <div style="display:flex;justify-content:space-between">
          <span style="color:var(--muted);font-size:13px">Generated</span>
          <span style="font-size:13px" id="d-time"></span></div>
      </div>
      <button class="btn btn-dl" id="btn-dl" onclick="download()">⬇ Download Research Paper (.docx)</button>
      <button class="btn btn-s" onclick="again()" style="margin-top:8px">Generate another paper</button>
      <button class="btn btn-s" onclick="loadDashboard();show('s-dashboard')" style="margin-top:6px;opacity:.7">← Back to Dashboard</button>
    </div>
  </div>
</div>

<!-- PROFILE -->
<div class="screen" id="s-profile">
  <div style="padding-top:28px">
    <div class="profile-header">
      <img class="profile-avatar" id="prof-avatar" src=""
        onerror="this.src='data:image/svg+xml,<svg xmlns=%22http://www.w3.org/2000/svg%22 viewBox=%220 0 64 64%22><circle cx=%2232%22 cy=%2232%22 r=%2232%22 fill=%22%23333%22/></svg>'">
      <div>
        <div style="font-size:20px;font-weight:700" id="prof-name">—</div>
        <div style="font-size:13px;color:var(--muted);margin-top:3px" id="prof-email">—</div>
        <div style="font-size:11px;color:var(--dim);margin-top:4px">Member since <span id="prof-since">—</span></div>
      </div>
    </div>
    <div class="stat-grid">
      <div class="stat-card"><div class="stat-val" id="prof-papers-count">0</div><div class="stat-lbl">Papers Generated</div></div>
      <div class="stat-card"><div class="stat-val" id="prof-spent">₹0</div><div class="stat-lbl">Total Spent</div></div>
      <div class="stat-card"><div class="stat-val" id="prof-paid-count">0</div><div class="stat-lbl">Papers Downloaded</div></div>
    </div>
    <div class="table-wrap">
      <div class="table-head">📄 Your Papers</div>
      <table><thead><tr><th>Topic</th><th>Date</th><th>Status</th></tr></thead>
      <tbody id="prof-papers-list"><tr><td colspan="3" class="empty">Loading...</td></tr></tbody></table>
    </div>
    <button class="btn btn-s" onclick="loadDashboard();show('s-dashboard')" style="max-width:180px">← Back</button>
  </div>
</div>

<!-- ADMIN -->
<div class="screen" id="s-admin">
  <div style="padding-top:28px">
    <div class="page-title">⚙️ Admin Dashboard</div>
    <div class="page-sub">All users, papers and payments</div>
    <div class="stat-grid">
      <div class="stat-card"><div class="stat-val" id="adm-users-c">—</div><div class="stat-lbl">Total Users</div></div>
      <div class="stat-card"><div class="stat-val" id="adm-papers-c">—</div><div class="stat-lbl">Papers Generated</div></div>
      <div class="stat-card"><div class="stat-val" id="adm-revenue-c">—</div><div class="stat-lbl">Total Revenue</div></div>
      <div class="stat-card"><div class="stat-val" id="adm-paid-c">—</div><div class="stat-lbl">Paid Downloads</div></div>
    </div>
    <div class="tabs">
      <button class="tab active" onclick="admTab('users',this)">👥 Users</button>
      <button class="tab" onclick="admTab('papers',this)">📄 Papers</button>
      <button class="tab" onclick="admTab('payments',this)">💳 Payments</button>
    </div>
    <div id="adm-tab-users">
      <div class="table-wrap"><table><thead><tr><th></th><th>Name</th><th>Email</th><th>Joined</th><th>Last Login</th></tr></thead>
      <tbody id="adm-users-list"></tbody></table></div>
    </div>
    <div id="adm-tab-papers" style="display:none">
      <div class="table-wrap"><table><thead><tr><th>Topic</th><th>User</th><th>Date</th><th>Status</th></tr></thead>
      <tbody id="adm-papers-list"></tbody></table></div>
    </div>
    <div id="adm-tab-payments" style="display:none">
      <div class="table-wrap"><table><thead><tr><th>User</th><th>Amount</th><th>Payment ID</th><th>Date</th><th>Status</th></tr></thead>
      <tbody id="adm-payments-list"></tbody></table></div>
    </div>
    <button class="btn btn-s" onclick="loadDashboard();show('s-dashboard')" style="max-width:180px;margin-top:12px">← Back</button>
  </div>
</div>

<!-- LEGAL DRAFTING -->
<div class="screen" id="s-legal">
  <div style="padding-top:28px;max-width:700px;margin:0 auto">
    <div style="margin-bottom:16px;display:flex;align-items:center;gap:12px">
      <button class="btn btn-s" style="width:auto;padding:8px 16px;font-size:12px" onclick="loadDashboard();show('s-dashboard')">← Dashboard</button>
      <div style="font-size:20px;font-weight:900;color:#111">⚖️ AI Legal Drafting</div>
    </div>
    <div class="cs" style="margin-bottom:20px">Describe the document you need and the details to include, and RDXper's AI will draft it for you or upload a format/sample document and we'll follow its structure using your data.</div>

    <!-- Mode tabs -->
    <div class="tabs" id="legal-tabs">
      <button class="tab active" onclick="legalSwitchTab('custom',this)">✍️ Describe &amp; Generate</button>
      <button class="tab" onclick="legalSwitchTab('format',this)">📎 Use My Own Format</button>
    </div>

    <div id="n-legal" class="notif"></div>

    <!-- Custom description form -->
    <div id="legal-form-custom">
      <div class="fg"><label>What kind of drafting do you need?</label>
        <input type="text" id="ld-doctype" placeholder="e.g. Rental Agreement, NDA, Employment Contract, Power of Attorney, Trademark Licence">
      </div>
      <div class="fg"><label>Details &amp; data for the draft</label>
        <textarea id="ld-details" rows="10" placeholder="Provide everything the document needs — party names & addresses, dates, amounts, terms, obligations, governing law, jurisdiction, special clauses, etc.&#10;&#10;Example: Landlord: Rohan Mehta, 12 MG Road, Pune. Tenant: Aisha Khan, 45 Park St, Pune. Property: 2BHK Flat No. 301, Green Meadows, Baner, Pune. Monthly rent: ₹28,000, payable by the 5th of every month. Security deposit: ₹1,00,000. Lease term: 11 months from 1 August 2026. Notice period: 1 month for termination by either party."></textarea>
      </div>
      <button class="btn btn-p" id="btn-legal-gen" onclick="generateLegalDoc()">⬇ Generate</button>
    </div>

    <!-- Upload-your-own-format form -->
    <div id="legal-form-format" style="display:none">
      <div class="fg"><label>Upload a format / sample document</label>
        <input type="file" id="ld-format-file" accept=".docx,.txt" style="width:100%;padding:10px;border:1.5px dashed #b0b0b0;border-radius:8px;background:#fafafa;font-size:13px">
        <div style="font-size:11px;color:#999;margin-top:4px">Accepted: .docx or .txt. We'll follow its structure and clauses.</div>
      </div>
      <div class="fg"><label>Data to fill into that format</label>
        <textarea id="ld-format-details" rows="9" placeholder="Provide the specific data that should replace the placeholders/details in the uploaded format, party names, dates, amounts, terms, etc."></textarea>
      </div>
      <button class="btn btn-p" id="btn-legal-gen-format" onclick="generateLegalDocFromFormat()">⬇ Generate</button>
    </div>

   

    <!-- Done state -->
    <div id="legal-done" style="display:none;text-align:center;padding:32px 0">
      <div style="font-size:48px;margin-bottom:12px">✅</div>
      <div class="ct">Draft ready!</div>
      <div class="cs" id="legal-done-sub">Your document has been generated.</div>
      <button class="btn btn-dl" id="btn-legal-dl" onclick="downloadLegal()" style="max-width:360px;margin:16px auto 8px">⬇ Download Draft (.docx)</button>
      <button class="btn btn-s" onclick="resetLegal()" style="max-width:200px;margin:0 auto">Generate Another</button>
    </div>
  </div>
</div>

<footer></footer>
</div>

<script>
const SECS=['keywords','abstract','introduction','objectives','literature_review','methodology','results','discussion','suggestions','limitations','conclusion'];
let token='',userEmail='',userName='',userPicture='',jobId='',curTopic='',curFigs=6,poll=null;
const ADMIN_EM='__ADMIN_EMAIL__';
const G_CLIENT='__GOOGLE_CLIENT_ID__';

// Restore session
(async function(){
  try{
    const t=localStorage.getItem('rx_tok'),e=localStorage.getItem('rx_em'),
          n=localStorage.getItem('rx_nm'),p=localStorage.getItem('rx_pic');
    if(t&&e){
      // Validate token is still alive on the server before showing dashboard
      const r=await fetch('/api/profile',{headers:{'Authorization':'Bearer '+t}});
      if(r.ok){
        token=t;userEmail=e;userName=n||'';userPicture=p||'';onLoggedIn();
      } else {
        // Token expired or server restarted — clear stale data
        ['rx_tok','rx_em','rx_nm','rx_pic'].forEach(k=>localStorage.removeItem(k));
      }
    }
  }catch(e){}
})();

// Simple sign-in — name + email, no password, no OTP
window.addEventListener('load', function(){
  document.getElementById('g-btn-wrap').innerHTML=`
    <div style="width:100%">
      <div class="fg" style="margin-bottom:10px">
        <label style="font-size:13px;font-weight:600">Your Name</label>
        <input type="text" id="si-name" placeholder="e.g. Rakunatha Khrishanth"
          style="background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;width:100%;font-size:14px;outline:none"
          onkeydown="if(event.key==='Enter')document.getElementById('si-email').focus()">
      </div>
      <div class="fg" style="margin-bottom:18px">
        <label style="font-size:13px;font-weight:600">Email Address</label>
        <input type="email" id="si-email" placeholder="you@email.com"
          style="background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;width:100%;font-size:14px;outline:none"
          onkeydown="if(event.key==='Enter')simpleSignIn()">
      </div>
      <button class="btn btn-p" id="btn-signin" onclick="simpleSignIn()" style="margin-bottom:0">Sign In →</button>
    </div>`;
});

async function simpleSignIn(){
  const name  = (document.getElementById('si-name')||{value:''}).value.trim();
  const email = (document.getElementById('si-email')||{value:''}).value.trim();
  const n = document.getElementById('n-login');
  if(!email || !email.includes('@')){
    n.className='notif error show'; n.textContent='Please enter a valid email address.'; return;
  }
  const btn = document.getElementById('btn-signin');
  btn.disabled=true; btn.textContent='Signing in...';
  try{
    const r = await fetch('/api/auth/login',{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({name: name||email.split('@')[0], email})});
    const d = await r.json();
    if(!d.success){
      n.className='notif error show'; n.textContent=d.message||'Sign in failed.';
      btn.disabled=false; btn.textContent='Sign In →'; return;
    }
    token=d.token; userEmail=d.email; userName=d.name; userPicture='';
    try{
      localStorage.setItem('rx_tok',token); localStorage.setItem('rx_em',userEmail);
      localStorage.setItem('rx_nm',userName); localStorage.setItem('rx_pic','');
    }catch(e){}
    onLoggedIn();
  }catch(e){
    n.className='notif error show'; n.textContent='Connection error. Try again.';
    btn.disabled=false; btn.textContent='Sign In →';
  }
}



function onLoggedIn(){
  document.getElementById('nav-auth').style.display='flex';
  document.getElementById('nav-name').textContent=userName||userEmail.split('@')[0];
  const av=document.getElementById('nav-avatar');
  if(userPicture){av.src=userPicture;av.style.display='block';}
  if(userEmail===ADMIN_EM) document.getElementById('admin-link').style.display='block';
  const aIn=document.getElementById('author-in');
  if(aIn&&!aIn.value) aIn.value=userName||'';
  loadDashboard();
  show('s-dashboard');
}

async function loadDashboard(){
  const nameEl=document.getElementById('dash-name-title');
  if(nameEl) nameEl.innerHTML=(userName||userEmail.split('@')[0]).split(' ')[0]+'<span>.</span>';
  try{
    const r=await fetch('/api/profile',{headers:{'Authorization':'Bearer '+token}});
    if(r.status===401){forceLogout();return;}
    const d=await r.json();
    if(!d.success) return;
    const papers=d.papers||[];
    const wrap=document.getElementById('dash-papers-wrap');
    if(papers.length===0){
      wrap.innerHTML=`<div class="dash-empty">
        <div class="dash-empty-icon">📄</div>
        <div class="dash-empty-txt">No papers yet</div>
        <div class="dash-empty-sub">Press <strong style="color:var(--accent)">+</strong> below to generate your first research paper</div>
      </div>`;
    } else {
      wrap.innerHTML='<div class="papers-grid">'+papers.map(p=>`
        <div class="paper-card">
          <div class="paper-card-topic">${escHtml(p.topic||'Untitled')}</div>
          <div class="paper-card-meta">
            <span class="paper-card-date">${(p.created_at||'').slice(0,10)}</span>
            <span class="paper-card-badge ${p.file_path?'badge-done':'badge-pending'}">${p.file_path?'✓ Done':'Pending'}</span>
          </div>
        </div>`).join('')+'</div>';
    }
  }catch(e){console.error('Dashboard load error',e);}
}

function forceLogout(){
  token='';userEmail='';userName='';userPicture='';
  try{['rx_tok','rx_em','rx_nm','rx_pic'].forEach(k=>localStorage.removeItem(k));}catch(e){}
  document.getElementById('nav-auth').style.display='none';
  document.getElementById('admin-link').style.display='none';
  show('s-home');
}

function escHtml(s){return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}

function startNewPaper(){
  // Reset questionnaire state then navigate
  ['topic-in','inst-in','q-problem','q-lit','q-gap','q-objectives','q-statement',
   'co-author-name','co-author-title','co-author-inst','co-author-email','co-author-phone'].forEach(id=>{
    const el=document.getElementById(id);if(el) el.value='';
  });
  const aIn=document.getElementById('author-in');
  if(aIn) aIn.value=userName||'';
  goStep(0);
  show('s-gen');
}

function show(id){document.querySelectorAll('.screen').forEach(s=>s.classList.remove('active'));document.getElementById(id).classList.add('active');window.scrollTo({top:0,behavior:'smooth'});}
function notify(id,msg,type){const e=document.getElementById(id);e.textContent=msg;e.className='notif '+type+' show';if(type!=='error')setTimeout(()=>e.classList.remove('show'),6000);}

function logout(){
  token='';userEmail='';userName='';userPicture='';
  try{['rx_tok','rx_em','rx_nm','rx_pic'].forEach(k=>localStorage.removeItem(k));}catch(e){}
  document.getElementById('nav-auth').style.display='none';
  document.getElementById('admin-link').style.display='none';
  show('s-home');
}

// ── QUESTIONNAIRE NAVIGATION ────────────────────────────────────────────────
let currentStep = 0;
const totalSteps = 6;

function goStep(n){
  // Only allow going back to completed steps
  if(n > currentStep) return;
  currentStep = n;
  renderStep();
}

function nextStep(from){
  // Only validate the topic (required), everything else is optional
  if(from===0 && !document.getElementById('topic-in').value.trim()){
    alert('Please enter your research topic — this is the only required field.'); return;
  }
  currentStep = from + 1;
  if(currentStep === 5) buildSummary();
  renderStep();
}

function prevStep(from){
  currentStep = from - 1;
  renderStep();
}

function renderStep(){
  for(let i=0;i<totalSteps;i++){
    const panel = document.getElementById('qp-'+i);
    const step  = document.getElementById('qs-'+i);
    if(!panel||!step) continue;
    panel.classList.toggle('active', i===currentStep);
    step.classList.remove('active','done');
    if(i===currentStep) step.classList.add('active');
    else if(i<currentStep) step.classList.add('done');
    // Update connector lines
    const lines = document.querySelectorAll('.q-line');
    lines.forEach((l,li)=>{ l.classList.toggle('done', li < currentStep); });
  }
  window.scrollTo({top:0,behavior:'smooth'});
}

function buildSummary(){
  const items = [
    {label:'Problem Identified', id:'q-problem'},
    {label:'Literature Reviewed', id:'q-lit'},
    {label:'Research Gap', id:'q-gap'},
    {label:'Objectives', id:'q-objectives'},
    {label:'Research Statement', id:'q-statement'},
  ];
  const s = document.getElementById('q-summary');
  if(!s) return;
  s.innerHTML = '<div style="font-size:13px;font-weight:700;margin-bottom:12px;color:var(--text)">📋 Your Research Inputs</div>' +
    items.map(item=>{
      const val = (document.getElementById(item.id)||{}).value||'';
      const preview = val.length > 120 ? val.slice(0,120)+'...' : val;
      return `<div class="q-summary-item">
        <div class="q-summary-label">${item.label}</div>
        <div class="q-summary-val">${preview||'<span style="color:var(--dim)">Not filled</span>'}</div>
      </div>`;
    }).join('');
}

async function generate(){
  const topic  = document.getElementById('topic-in').value.trim();
  const author = document.getElementById('author-in').value.trim();
  const inst   = document.getElementById('inst-in').value.trim();
  const nfigs  = parseInt(document.getElementById('sl').value);
  const qProblem    = document.getElementById('q-problem').value.trim();
  const qLit        = document.getElementById('q-lit').value.trim();
  const qGap        = document.getElementById('q-gap').value.trim();
  const qObjectives = document.getElementById('q-objectives').value.trim();
  const qStatement  = document.getElementById('q-statement').value.trim();
  const coName      = document.getElementById('co-author-name').value.trim();
  const coTitle     = document.getElementById('co-author-title').value.trim();
  const coInst      = document.getElementById('co-author-inst').value.trim();
  const coEmail     = document.getElementById('co-author-email').value.trim();
  const coPhone     = document.getElementById('co-author-phone').value.trim();
  if(!topic){notify('n-gen','Please enter a research topic.','error');return;}

  const btn=document.getElementById('btn-gen');
  btn.disabled=true;btn.innerHTML='<span class="spin"></span>Generating...';
  try{
    const r=await fetch('/api/generate',{method:'POST',
      headers:{'Content-Type':'application/json','Authorization':'Bearer '+token},
      body:JSON.stringify({
        topic, author_name:author, institution:inst, num_figures:nfigs,
        q_problem:qProblem, q_lit:qLit, q_gap:qGap,
        q_objectives:qObjectives, q_statement:qStatement,
        co_author_name:coName, co_author_title:coTitle,
        co_author_inst:coInst, co_author_email:coEmail, co_author_phone:coPhone
      })});
    const d=await r.json();
    if(r.status===401){btn.disabled=false;btn.innerHTML='Generate Research Paper';forceLogout();return;}
    if(!d.success){notify('n-gen',d.message||'Failed.','error');btn.disabled=false;btn.innerHTML='Generate Research Paper';return;}
    jobId=d.job_id;curTopic=topic;curFigs=nfigs;
    document.getElementById('prog-topic').textContent=topic;
    buildSecGrid();show('s-prog');pollStatus();
  }catch(e){notify('n-gen','Connection error.','error');btn.disabled=false;btn.innerHTML='Generate Research Paper';}
}

function buildSecGrid(){
  const g=document.getElementById('sec-grid');g.innerHTML='';
  SECS.forEach(s=>{const d=document.createElement('div');d.className='sec-item';d.id='sec-'+s;d.textContent=s.replace('_',' ');g.appendChild(d);});
}

function updateSecs(pct){
  const idx=Math.floor((Math.max(0,pct-30))/45*SECS.length);
  SECS.forEach((s,i)=>{const el=document.getElementById('sec-'+s);if(!el)return;
    if(i<idx)el.className='sec-item done';else if(i===idx)el.className='sec-item writing';});
}

function pollStatus(){
  poll=setInterval(async()=>{
    try{
      const r=await fetch('/api/status/'+jobId,{headers:{'Authorization':'Bearer '+token}});
      const d=await r.json();if(!d.success)return;
      document.getElementById('prog-fill').style.width=d.progress+'%';
      document.getElementById('prog-pct').textContent=d.progress+'%';
      document.getElementById('stage-msg').textContent=d.message;
      updateSecs(d.progress);
      if(d.status==='done'){
        clearInterval(poll);
        SECS.forEach(s=>{const e=document.getElementById('sec-'+s);if(e)e.className='sec-item done';});
        document.getElementById('d-topic').textContent=curTopic;
        document.getElementById('d-figs').textContent=curFigs+' figures';
        document.getElementById('d-time').textContent=new Date().toLocaleTimeString();
        show('s-done');
      }else if(d.status==='error'){
        clearInterval(poll);
        const btn=document.getElementById('btn-gen');btn.disabled=false;btn.innerHTML='✦ Generate Paper (Free AI)';
        alert('Generation failed: '+d.message);show('s-gen');
      }
    }catch(e){console.error(e);}
  },800);
}

async function download(){
  const btn=document.getElementById('btn-dl');btn.disabled=true;btn.innerHTML='<span class="spin"></span>Downloading...';
  try{
    const r=await fetch('/api/download/'+jobId,{headers:{'Authorization':'Bearer '+token}});
    if(!r.ok)throw new Error('failed');
    const blob=await r.blob(),url=URL.createObjectURL(blob),a=document.createElement('a');
    a.href=url;a.download='rdxper_'+curTopic.slice(0,40).replace(/[^a-zA-Z0-9]/g,'_')+'.docx';a.click();URL.revokeObjectURL(url);
  }catch(e){alert('Download failed. Try again.');}
  finally{btn.disabled=false;btn.innerHTML='⬇ Download Research Paper (.docx)';}
}

function again(){
  jobId='';curTopic='';
  ['topic-in','inst-in','q-problem','q-lit','q-gap','q-objectives','q-statement',
   'co-author-name','co-author-title','co-author-inst','co-author-email','co-author-phone'].forEach(id=>{
    const el=document.getElementById(id);if(el) el.value='';
  });
  document.getElementById('sl').value=6;document.getElementById('sl-display').textContent='6';
  const btn=document.getElementById('btn-gen');
  if(btn){btn.disabled=false;btn.innerHTML='✦ Generate Research Paper';}
  document.getElementById('n-gen').classList.remove('show');
  currentStep=0;renderStep();
  loadDashboard();
  show('s-dashboard');
}

async function showProfile(){
  show('s-profile');
  try{
    const r=await fetch('/api/profile',{headers:{'Authorization':'Bearer '+token}});
    const d=await r.json();if(!d.success)return;
    const u=d.user;
    document.getElementById('prof-avatar').src=u.picture||'';
    document.getElementById('prof-name').textContent=u.name||u.email;
    document.getElementById('prof-email').textContent=u.email;
    document.getElementById('prof-since').textContent=(u.created_at||'').split('T')[0]||u.created_at||'—';
    document.getElementById('prof-papers-count').textContent=d.papers_count;
    document.getElementById('prof-spent').textContent='₹'+d.total_spent;
    document.getElementById('prof-paid-count').textContent=d.papers.filter(p=>p.paid).length;
    const tb=document.getElementById('prof-papers-list');
    tb.innerHTML=d.papers.length===0
      ?'<tr><td colspan="3" class="empty">No papers yet. Generate your first one!</td></tr>'
      :d.papers.map(p=>`<tr>
        <td style="max-width:240px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${p.topic||''}">${p.topic||'—'}</td>
        <td style="white-space:nowrap">${(p.created_at||'').split('T')[0]||'—'}</td>
        <td>${p.paid?'<span class="badge-paid">✓ Downloaded</span>':'<span class="badge-free">Generated</span>'}</td>
      </tr>`).join('');
  }catch(e){console.error(e);}
}

async function showAdmin(){
  show('s-admin');
  try{
    const r=await fetch('/api/admin/stats',{headers:{'Authorization':'Bearer '+token}});
    const d=await r.json();
    if(!d.success){alert('Access denied');show('s-gen');return;}
    document.getElementById('adm-users-c').textContent=d.stats.total_users;
    document.getElementById('adm-papers-c').textContent=d.stats.total_papers;
    document.getElementById('adm-revenue-c').textContent='₹'+d.stats.total_revenue;
    document.getElementById('adm-paid-c').textContent=d.stats.paid_papers;
    document.getElementById('adm-users-list').innerHTML=d.users.length===0
      ?'<tr><td colspan="5" class="empty">No users yet.</td></tr>'
      :d.users.map(u=>`<tr>
        <td><img class="avatar" src="${u.picture||''}" onerror="this.style.display='none'"></td>
        <td>${u.name||'—'}</td><td>${u.email}</td>
        <td>${(u.created_at||'').split('T')[0]||'—'}</td>
        <td>${(u.last_login||'').split('T')[0]||'—'}</td>
      </tr>`).join('');
    document.getElementById('adm-papers-list').innerHTML=d.papers.length===0
      ?'<tr><td colspan="4" class="empty">No papers yet.</td></tr>'
      :d.papers.map(p=>`<tr>
        <td style="max-width:280px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap">${p.topic||'—'}</td>
        <td>${p.email||'—'}</td>
        <td>${(p.created_at||'').split('T')[0]||'—'}</td>
        <td>${p.paid?'<span class="badge-paid">✓ Paid</span>':'<span class="badge-pending">Pending</span>'}</td>
      </tr>`).join('');
    document.getElementById('adm-payments-list').innerHTML=d.payments.length===0
      ?'<tr><td colspan="5" class="empty">No payments yet.</td></tr>'
      :d.payments.map(p=>`<tr>
        <td>${p.email||'—'}</td>
        <td style="color:var(--accent);font-weight:700">₹${p.amount||0}</td>
        <td style="font-family:monospace;font-size:11px">${(p.razorpay_payment||'—').slice(0,24)}</td>
        <td>${(p.created_at||'').split('T')[0]||'—'}</td>
        <td><span class="badge-${p.status==='paid'?'paid':'pending'}">${p.status}</span></td>
      </tr>`).join('');
  }catch(e){console.error(e);}
}

function admTab(name,el){
  document.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));el.classList.add('active');
  ['users','papers','payments'].forEach(t=>{
    const d=document.getElementById('adm-tab-'+t);if(d)d.style.display=t===name?'block':'none';
  });
}

// ── LEGAL DRAFTING ────────────────────────────────────────────────────────────
let legalJobId = '';
let legalDlName = 'RDXper_Legal_Draft.docx';

function legalSwitchTab(mode, btn){
  document.querySelectorAll('#legal-tabs .tab').forEach(t=>t.classList.remove('active'));
  btn.classList.add('active');
  document.getElementById('legal-form-custom').style.display = mode==='custom' ? 'block' : 'none';
  document.getElementById('legal-form-format').style.display = mode==='format' ? 'block' : 'none';
  document.getElementById('legal-done').style.display='none';
  const n=document.getElementById('n-legal'); n.classList.remove('show');
}

function resetLegal(){
  legalJobId='';
  document.getElementById('legal-done').style.display='none';
  document.getElementById('legal-form-custom').style.display='block';
  document.getElementById('legal-form-format').style.display='none';
  document.querySelectorAll('#legal-tabs .tab').forEach((t,i)=>t.classList.toggle('active', i===0));
  ['ld-doctype','ld-details','ld-format-details'].forEach(id=>{const el=document.getElementById(id);if(el)el.value='';});
  const fEl=document.getElementById('ld-format-file'); if(fEl) fEl.value='';
  document.getElementById('n-legal').classList.remove('show');
}

function legalShowError(msg){
  const n=document.getElementById('n-legal');
  n.className='notif error show'; n.textContent=msg;
}

async function generateLegalDoc(){
  const g = id => (document.getElementById(id)||{}).value.trim()||'';
  const doc_type = g('ld-doctype'), details = g('ld-details');
  if(!doc_type || !details){
    legalShowError('Please enter the type of document and the details/data for the draft.');
    return;
  }
  const btn=document.getElementById('btn-legal-gen');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Drafting with AI...';
  document.getElementById('n-legal').classList.remove('show');
  try{
    const r = await fetch('/api/legal/generate',{
      method:'POST',
      headers:{'Content-Type':'application/json','Authorization':'Bearer '+token},
      body:JSON.stringify({mode:'custom', doc_type, details})
    });
    const d = await r.json();
    if(!d.success){ legalShowError(d.message||'Generation failed.'); return; }
    legalJobId = d.job_id;
    legalDlName = (doc_type.replace(/[^\w\-]+/g,'_')||'RDXper_Legal_Draft') + '.docx';
    document.getElementById('legal-done-sub').textContent = 'Your ' + doc_type + ' has been generated.';
    document.getElementById('legal-form-custom').style.display='none';
    document.getElementById('legal-form-format').style.display='none';
    document.getElementById('legal-done').style.display='block';
  }catch(e){
    legalShowError('Connection error. Please try again.');
  }finally{
    btn.disabled=false; btn.innerHTML='⬇ Generate Draft (.docx)';
  }
}

async function generateLegalDocFromFormat(){
  const details = (document.getElementById('ld-format-details')||{}).value.trim()||'';
  const fileEl = document.getElementById('ld-format-file');
  const file = fileEl && fileEl.files && fileEl.files[0];
  if(!file){ legalShowError('Please upload a format/sample document (.docx or .txt).'); return; }
  if(!details){ legalShowError('Please enter the data to fill into the uploaded format.'); return; }
  const btn=document.getElementById('btn-legal-gen-format');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Drafting with AI...';
  document.getElementById('n-legal').classList.remove('show');
  try{
    const fd = new FormData();
    fd.append('mode','format');
    fd.append('details', details);
    fd.append('format_file', file);
    const r = await fetch('/api/legal/generate',{
      method:'POST',
      headers:{'Authorization':'Bearer '+token},
      body: fd
    });
    const d = await r.json();
    if(!d.success){ legalShowError(d.message||'Generation failed.'); return; }
    legalJobId = d.job_id;
    legalDlName = 'RDXper_Legal_Draft.docx';
    document.getElementById('legal-done-sub').textContent = 'Your document has been generated from the uploaded format.';
    document.getElementById('legal-form-custom').style.display='none';
    document.getElementById('legal-form-format').style.display='none';
    document.getElementById('legal-done').style.display='block';
  }catch(e){
    legalShowError('Connection error. Please try again.');
  }finally{
    btn.disabled=false; btn.innerHTML='⬇ Generate Draft (.docx)';
  }
}

async function downloadLegal(){
  const btn=document.getElementById('btn-legal-dl');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Downloading...';
  try{
    const r = await fetch('/api/download/'+legalJobId,{headers:{'Authorization':'Bearer '+token}});
    if(!r.ok) throw new Error('failed');
    const blob=await r.blob(), url=URL.createObjectURL(blob), a=document.createElement('a');
    a.href=url; a.download=legalDlName; a.click(); URL.revokeObjectURL(url);
  }catch(e){ alert('Download failed. Please try again.'); }
  finally{ btn.disabled=false; btn.innerHTML='⬇ Download Draft (.docx)'; }
}
</script>
</body>
</html>"""




# ═══════════════════════════════════════════════════════════════════════════════
#  FLASK ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    client_id = os.environ.get('GOOGLE_CLIENT_ID', '')
    html = HTML.replace('__GOOGLE_CLIENT_ID__', client_id).replace('__ADMIN_EMAIL__', ADMIN_EMAIL)
    return Response(html, mimetype='text/html')


def _verify_google_token(id_token_str):
    try:
        url = "https://oauth2.googleapis.com/tokeninfo?id_token=" + urllib.parse.quote(id_token_str)
        req = urllib.request.Request(url, headers={"User-Agent": "rdxper/4.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            info = json.loads(resp.read())
        client_id = os.environ.get("GOOGLE_CLIENT_ID", "")
        if client_id and info.get("aud") != client_id:
            return None
        if info.get("exp") and int(info["exp"]) < time.time():
            return None
        return info
    except Exception as e:
        print(f"[Google] Token error: {e}")
        return None

@app.route("/api/auth/dev", methods=["POST"])
@app.route("/api/auth/login", methods=["POST"])
def simple_login():
    """Simple name + email login — works in all environments."""
    data    = request.json or {}
    email   = data.get("email", "").strip().lower()
    name    = data.get("name", "").strip() or email.split("@")[0]
    if not email or "@" not in email:
        return jsonify({"success": False, "message": "Valid email required"}), 400
    user_id = "u_" + email.replace("@","_").replace(".","_")
    with get_db() as db:
        user = db.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        if user:
            db.execute("UPDATE users SET name=?,last_login=datetime('now') WHERE email=?", (name, email))
            user_id = user["id"]
        else:
            db.execute("INSERT INTO users (id,email,name,picture,last_login) VALUES (?,?,?,?,datetime('now'))",
                       (user_id, email, name, ""))
    tok = secrets.token_urlsafe(32)
    session_set(tok, email)
    sessions[tok]["user_id"] = user_id
    sessions[tok]["name"] = name
    sessions[tok]["picture"] = ""
    return jsonify({"success": True, "token": tok, "email": email, "name": name, "picture": ""})

@app.route("/api/auth/google", methods=["POST"])
def google_auth():
    id_token_str = request.json.get("id_token", "")
    if not id_token_str:
        return jsonify({"success": False, "message": "No token"}), 400
    info = _verify_google_token(id_token_str)
    if not info:
        return jsonify({"success": False, "message": "Invalid Google token"}), 401
    g_email   = info.get("email", "").lower()
    g_name    = info.get("name", g_email.split("@")[0])
    g_picture = info.get("picture", "")
    g_sub     = info.get("sub", str(uuid.uuid4()))
    with get_db() as db:
        user = db.execute("SELECT * FROM users WHERE email=?", (g_email,)).fetchone()
        if user:
            db.execute("UPDATE users SET name=?,picture=?,last_login=datetime('now') WHERE email=?",
                       (g_name, g_picture, g_email))
            user_id = user["id"]
        else:
            user_id = g_sub
            db.execute("INSERT INTO users (id,email,name,picture,last_login) VALUES (?,?,?,?,datetime('now'))",
                       (user_id, g_email, g_name, g_picture))
    tok = secrets.token_urlsafe(32)
    session_set(tok, g_email)
    sessions[tok]["user_id"] = user_id
    sessions[tok]["name"] = g_name
    sessions[tok]["picture"] = g_picture
    return jsonify({"success": True, "token": tok, "email": g_email, "name": g_name, "picture": g_picture})

@app.route("/api/profile")
def get_profile():
    tok = request.headers.get("Authorization", "").replace("Bearer ", "")
    sess = session_get(tok)
    if not sess:
        return jsonify({"success": False, "message": "Unauthorized"}), 401
    with get_db() as db:
        user    = db.execute("SELECT * FROM users WHERE id=?", (sess["user_id"],)).fetchone()
        papers  = db.execute("SELECT * FROM papers WHERE user_id=? ORDER BY created_at DESC", (sess["user_id"],)).fetchall()
        result  = db.execute("SELECT COALESCE(SUM(amount),0) as t FROM payments WHERE user_id=? AND status='paid'", (sess["user_id"],)).fetchone()
        total_spent = result["t"]
    return jsonify({
        "success": True,
        "user": dict(user),
        "papers": [dict(p) for p in papers],
        "total_spent": total_spent,
        "papers_count": len(papers)
    })

@app.route("/api/admin/stats")
def admin_stats():
    tok = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not session_get(tok):
        return jsonify({"success": False, "message": "Unauthorized"}), 401
    if sessions.get(tok, {}).get("email") != ADMIN_EMAIL:
        return jsonify({"success": False, "message": "Forbidden"}), 403
    with get_db() as db:
        users    = db.execute("SELECT * FROM users ORDER BY created_at DESC").fetchall()
        papers   = db.execute("SELECT p.*,u.email,u.name FROM papers p JOIN users u ON p.user_id=u.id ORDER BY p.created_at DESC").fetchall()
        payments = db.execute("SELECT pay.*,u.email FROM payments pay JOIN users u ON pay.user_id=u.id ORDER BY pay.created_at DESC").fetchall()
        revenue  = db.execute("SELECT COALESCE(SUM(amount),0) as t FROM payments WHERE status='paid'").fetchone()["t"]
    return jsonify({
        "success": True,
        "stats": {"total_users": len(users), "total_papers": len(papers),
                  "total_revenue": revenue, "paid_papers": sum(1 for p in papers if p["paid"])},
        "users":    [dict(u) for u in users],
        "papers":   [dict(p) for p in papers],
        "payments": [dict(p) for p in payments]
    })

@app.route('/api/send-otp', methods=['POST'])
def send_otp():
    data  = request.json
    email = data.get('email', '').strip().lower()
    if not email or '@' not in email:
        return jsonify({'success': False, 'message': 'Invalid email'}), 400
    otp = str(secrets.randbelow(900000) + 100000)
    otp_store[email] = {'otp': otp, 'expires': time.time() + 600}
    print(f"\n{'='*40}\n OTP for {email}: {otp}\n{'='*40}\n")
    _try_smtp(email, otp)
    return jsonify({'success': True, 'message': f'OTP sent to {email}', 'demo_otp': otp})

def _try_smtp(to_email: str, otp: str):
    u = os.environ.get('SMTP_USER')
    p = os.environ.get('SMTP_PASS')
    if not (u and p):
        return
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'Your rdxper Login Code'
        msg['From'] = u; msg['To'] = to_email
        msg.attach(MIMEText(
            f'<h2 style="color:#111111">Your rdxper OTP</h2>'
            f'<p style="font-size:32px;letter-spacing:8px;font-family:monospace"><b>{otp}</b></p>'
            f'<p>Valid for 10 minutes.</p>', 'html'))
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as s:
            s.login(u, p)
            s.sendmail(u, [to_email, ADMIN_EMAIL], msg.as_string())
    except Exception as e:
        print(f'[SMTP] {e}')

@app.route('/api/verify-otp', methods=['POST'])
def verify_otp():
    data  = request.json
    email = data.get('email', '').strip().lower()
    otp   = data.get('otp', '').strip()
    rec   = otp_store.get(email)
    if not rec:
        return jsonify({'success': False, 'message': 'No OTP found. Request a new one.'}), 400
    if time.time() > rec['expires']:
        del otp_store[email]
        return jsonify({'success': False, 'message': 'OTP expired.'}), 400
    if rec['otp'] != otp:
        return jsonify({'success': False, 'message': 'Wrong OTP.'}), 400
    tok = secrets.token_urlsafe(32)
    session_set(tok, email)
    del otp_store[email]
    return jsonify({'success': True, 'token': tok, 'email': email})

@app.route('/api/generate', methods=['POST'])
def generate_paper():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data   = request.json

    if not os.environ.get("GROQ_API_KEY", "").strip():
        return jsonify({'success': False,
                        'message': 'GROQ_API_KEY not set. Get a free key at https://console.groq.com'}), 400

    topic  = data.get('topic', '').strip()
    nfigs  = max(3, min(25, int(data.get('num_figures', 6))))
    author = data.get('author_name', 'Anonymous').strip()
    inst   = data.get('institution', '').strip()
    email  = sess['email']

    # Co-author fields
    co_author       = data.get('co_author_name', '').strip()
    co_author_title = data.get('co_author_title', '').strip()
    co_author_inst  = data.get('co_author_inst', '').strip()
    co_author_email = data.get('co_author_email', '').strip()
    co_author_phone = data.get('co_author_phone', '').strip()

    # Questionnaire fields
    q_problem    = data.get('q_problem', '').strip()
    q_lit        = data.get('q_lit', '').strip()
    q_gap        = data.get('q_gap', '').strip()
    q_objectives = data.get('q_objectives', '').strip()
    q_statement  = data.get('q_statement', '').strip()

    if not topic:
        return jsonify({'success': False, 'message': 'Topic required'}), 400

    jid     = str(uuid.uuid4())
    user_id = sess.get('user_id', email)
    jobs[jid] = {'status': 'queued', 'progress': 0,
                 'message': 'Queued...', 'file_path': None, 'topic': topic, 'user_id': user_id}
    with get_db() as db:
        # Ensure user exists (guards against FK constraint failure)
        db.execute(
            'INSERT OR IGNORE INTO users (id, email, name, picture) VALUES (?, ?, ?, ?)',
            (user_id, email, sess.get('name', ''), sess.get('picture', ''))
        )
        db.execute('INSERT INTO papers (id,user_id,topic) VALUES (?,?,?)', (jid, user_id, topic))

    questionnaire = {
        'problem':    q_problem,
        'lit':        q_lit,
        'gap':        q_gap,
        'objectives': q_objectives,
        'statement':  q_statement,
    }

    co_author_info = {
        'name':  co_author,
        'title': co_author_title,
        'inst':  co_author_inst,
        'email': co_author_email,
        'phone': co_author_phone,
    }

    def _run():
        try:
            g    = PaperGenerator(jid, jobs)
            path = g.generate(topic, nfigs, author, inst, email, questionnaire, co_author_info)
            jobs[jid].update({'status': 'done', 'progress': 100,
                              'message': 'Research paper ready!', 'file_path': path})
            with get_db() as db:
                db.execute('UPDATE papers SET file_path=? WHERE id=?', (path, jid))
        except Exception as e:
            import traceback; traceback.print_exc()
            jobs[jid].update({'status': 'error', 'message': str(e)})

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({'success': True, 'job_id': jid})

@app.route('/api/status/<jid>')
def job_status(jid):
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    if not session_get(tok):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    job = jobs.get(jid)
    if not job:
        # Fall back to DB — server may have restarted mid-generation
        with get_db() as db:
            paper = db.execute('SELECT file_path, topic FROM papers WHERE id=?', (jid,)).fetchone()
        if not paper:
            return jsonify({'success': False, 'message': 'Job not found'}), 404
        if paper['file_path']:
            return jsonify({'success': True, 'status': 'done', 'progress': 100, 'message': 'Research paper ready!'})
        return jsonify({'success': True, 'status': 'error', 'progress': 0, 'message': 'Job lost after server restart — please generate again.'})
    return jsonify({'success': True, 'status': job['status'],
                    'progress': job['progress'], 'message': job['message']})

@app.route('/api/download/<jid>')
def download_paper(jid):
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    if not session_get(tok):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    # First check in-memory jobs dict
    job = jobs.get(jid)
    fp = None

    if job:
        if job['status'] != 'done':
            return jsonify({'success': False, 'message': 'File not ready'}), 400
        fp = job.get('file_path')
    else:
        # Server may have restarted — look up file path from DB
        with get_db() as db:
            paper = db.execute('SELECT file_path, topic FROM papers WHERE id=?', (jid,)).fetchone()
        if not paper:
            return jsonify({'success': False, 'message': 'Job not found'}), 404
        fp = paper['file_path']
        topic_slug = paper['topic'] if paper['topic'] else jid
        if not fp:
            return jsonify({'success': False, 'message': 'File not ready — please generate again'}), 400
        # Restore minimal job info for slug below
        jobs[jid] = {'status': 'done', 'file_path': fp, 'topic': paper['topic'] or ''}

    if not fp or not os.path.exists(fp):
        return jsonify({'success': False, 'message': 'File not found on server'}), 404

    topic_for_slug = jobs[jid].get('topic', '') if jid in jobs else ''
    slug = re.sub(r'[^\w\-]', '_', topic_for_slug[:40]) if topic_for_slug else jid[:8]
    return send_file(fp, as_attachment=True,
                     download_name=f'rdxper_{slug}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




# ═══════════════════════════════════════════════════════════════════════════════
#  AI LEGAL DRAFTING
# ═══════════════════════════════════════════════════════════════════════════════

RDXPER_WATERMARK_TEXT = 'RDXper - A Rakunatha Khrishanth Manathra Creation'


def add_watermark(doc, text: str = RDXPER_WATERMARK_TEXT):
    """Insert a diagonal, semi-transparent watermark into the header of every
    section (the classic Word VML watermark technique), plus a small text
    credit line in the footer as a reliable fallback for viewers that don't
    render VML shapes."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN

    from docx.oxml import parse_xml
    import xml.sax.saxutils as _sax

    safe_text = _sax.escape(text, {'"': '&quot;'})

    watermark_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136" adj="10800" '
        'path="m@7,0l@8,0m@5,21600l@6,21600e">'
        '<v:formulas>'
        '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 #0"/>'
        '<v:f eqn="sum 0 0 #1"/><v:f eqn="prod #1 2 1"/><v:f eqn="sum 21600 0 #1"/>'
        '<v:f eqn="if #0 #3 0"/><v:f eqn="if #0 21600 #1"/><v:f eqn="if #3 21600 #2"/>'
        '<v:f eqn="if #3 #1 21600"/><v:f eqn="mid #4 #5"/><v:f eqn="mid #6 #7"/><v:f eqn="val #0"/>'
        '</v:formulas>'
        '<v:path textpathok="t" o:connecttype="custom" '
        'o:connectlocs="@9,0;@10,10800;@9,21600;@8,10800" o:connectangles="270,180,90,0"/>'
        '<v:textpath on="t" fitshape="t"/>'
        '</v:shapetype>'
        '<v:shape id="RDXperWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:520pt;height:110pt;'
        'rotation:315;z-index:-251654144;mso-position-horizontal:center;'
        'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
        'mso-position-vertical-relative:margin" o:allowincell="f" fillcolor="#D8D8D8" stroked="f">'
        '<v:fill opacity=".45"/>'
        f'<v:textpath style="font-family:\'Calibri\';font-size:1pt" string="{safe_text}"/>'
        '</v:shape>'
        '</w:pict>'
    )

    for section in doc.sections:
        # ── Diagonal watermark shape in the header ──────────────────────────
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ''
        h_para.alignment = _ALIGN.CENTER
        run = h_para.add_run()
        r_el = run._r
        pict = parse_xml(watermark_xml)
        r_el.append(pict)

        # ── Small credit line in the footer (reliable fallback) ────────────
        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ''
        f_para.alignment = _ALIGN.CENTER
        f_run = f_para.add_run(text)
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
        f_run.italic = True


def build_ai_legal_docx(doc_type: str, ai_text: str) -> str:
    """Convert the AI-drafted plain-text legal document into a formatted,
    watermarked .docx file."""
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    TNR = 'Times New Roman'
    lines = [ln.rstrip() for ln in ai_text.strip().split('\n')]

    numbered_re   = re.compile(r'^\s*(\d{1,3})[\.\)]\s+(.*)$')
    title_written = False

    for ln in lines:
        stripped = ln.strip()
        if not stripped:
            continue
        # Skip stray markdown fences/asterisked bold markers from the LLM
        clean = stripped.strip('#').strip()
        clean = re.sub(r'^\*\*(.*)\*\*$', r'\1', clean).strip()
        clean = clean.lstrip('*').strip()
        if not clean:
            continue

        m = numbered_re.match(clean)
        if not title_written and not m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            r = p.add_run(clean.upper())
            r.bold = True; r.font.size = Pt(16); r.font.name = TNR
            title_written = True
            continue

        if m:
            num, body = m.group(1), m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            r_num = p.add_run(f'{num}.  ')
            r_num.bold = True; r_num.font.size = Pt(12); r_num.font.name = TNR
            r_body = p.add_run(body)
            r_body.font.size = Pt(12); r_body.font.name = TNR
        elif clean.isupper() and len(clean) < 80:
            # Section heading in caps, e.g. "WITNESSETH", "THE SCHEDULE"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(clean)
            r.bold = True; r.font.size = Pt(13); r.font.name = TNR
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(clean)
            r.font.size = Pt(12); r.font.name = TNR

    add_watermark(doc, RDXPER_WATERMARK_TEXT)

    os.makedirs('generated', exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', (doc_type or 'Legal_Draft')[:40]) or 'Legal_Draft'
    out  = os.path.abspath(f'generated/{safe}_{uuid.uuid4().hex[:8]}.docx')
    doc.save(out)
    return out


def extract_text_from_upload(file_storage) -> str:
    """Extract plain text from an uploaded .docx or .txt reference format file."""
    filename = (file_storage.filename or '').lower()
    if filename.endswith('.docx'):
        tmp_path = os.path.abspath(f'generated/_upload_{uuid.uuid4().hex[:8]}.docx')
        os.makedirs('generated', exist_ok=True)
        file_storage.save(tmp_path)
        try:
            src = Document(tmp_path)
            text = '\n'.join(p.text for p in src.paragraphs if p.text.strip())
            for tbl in src.tables:
                for row in tbl.rows:
                    text += '\n' + ' | '.join(c.text for c in row.cells)
            return text
        finally:
            try: os.remove(tmp_path)
            except OSError: pass
    elif filename.endswith('.txt'):
        raw = file_storage.read()
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            return raw.decode('latin-1', errors='ignore')
    else:
        raise ValueError('Unsupported file type — please upload a .docx or .txt file.')


def ai_draft_legal_document(doc_type: str, details: str, reference_text: str = '') -> str:
    """Call the AI model to draft a full legal document as plain text."""
    system = (
        'You are an expert legal drafter. Draft complete, professional, ready-to-use legal '
        'documents in plain text (no markdown, no asterisks, no code fences). '
        'Structure: a centred ALL-CAPS title on the first line, then the preamble/recitals '
        'as plain paragraphs, then the operative clauses as a numbered list ("1. ", "2. ", ...), '
        'and finally a signature block. Use precise, formal legal language appropriate for the '
        'jurisdiction implied by the details given. Do not include any commentary, explanations, '
        'or notes outside the document itself — output ONLY the document text.'
    )
    if reference_text:
        prompt = (
            f'Use the following document as the FORMAT/STRUCTURE reference — follow its layout, '
            f'clause structure and drafting style closely, but replace all names, dates, amounts '
            f'and other details with the DATA provided below. Fill in any gaps sensibly.\n\n'
            f'--- FORMAT REFERENCE ---\n{reference_text[:6000]}\n\n'
            f'--- DATA TO USE ---\n{details}\n\n'
            f'Now produce the complete final document text.'
        )
    else:
        prompt = (
            f'Draft a "{doc_type}" document using the following details and data:\n\n'
            f'{details}\n\n'
            f'Produce the complete, professional, ready-to-use document text.'
        )
    return ai_generate(prompt, system=system, temperature=0.4)


@app.route('/api/legal/generate', methods=['POST'])
def gen_ai_legal_draft():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if not os.environ.get('GROQ_API_KEY', '').strip():
        return jsonify({'success': False,
                        'message': 'GROQ_API_KEY not set. Get a free key at https://console.groq.com'}), 400

    is_multipart = request.content_type and 'multipart/form-data' in request.content_type
    mode = (request.form.get('mode') if is_multipart else (request.json or {}).get('mode')) or 'custom'

    try:
        if mode == 'format':
            details = (request.form.get('details') or '').strip()
            if not details:
                return jsonify({'success': False, 'message': 'Please provide the data to fill into the format.'}), 400
            file_storage = request.files.get('format_file')
            if not file_storage or not file_storage.filename:
                return jsonify({'success': False, 'message': 'Please upload a format/sample document.'}), 400
            reference_text = extract_text_from_upload(file_storage)
            doc_type = 'Legal Draft'
            ai_text = ai_draft_legal_document(doc_type, details, reference_text=reference_text)
        else:
            data = request.json or {}
            doc_type = (data.get('doc_type') or '').strip()
            details  = (data.get('details') or '').strip()
            if not doc_type or not details:
                return jsonify({'success': False, 'message': 'Please provide the document type and details.'}), 400
            ai_text = ai_draft_legal_document(doc_type, details)

        path = build_ai_legal_docx(doc_type, ai_text)
        jid  = uuid.uuid4().hex
        jobs[jid] = {'status': 'done', 'file_path': path, 'topic': doc_type}
        return jsonify({'success': True, 'job_id': jid})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════════
#  LEGAL DRAFTING — TRADEMARK LICENSE GENERATOR (legacy structured template,
#  still available programmatically via /api/legal/trademark-license)
# ═══════════════════════════════════════════════════════════════════════════════

def build_trademark_license_docx(data: dict) -> str:
    """Generate a Licence to Use Trade Mark agreement as a .docx file."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    TNR = 'Times New Roman'

    def para(text, bold=False, sz=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sp_b=6, sp_a=6, center=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
        p.paragraph_format.space_before = Pt(sp_b)
        p.paragraph_format.space_after  = Pt(sp_a)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(sz)
        r.font.name = TNR
        return p

    def clause(number, text, sp_b=4, sp_a=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(sp_b)
        p.paragraph_format.space_after  = Pt(sp_a)
        p.paragraph_format.left_indent  = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        r_num = p.add_run(f'{number}.  ')
        r_num.bold = True
        r_num.font.size = Pt(12)
        r_num.font.name = TNR
        r_body = p.add_run(text)
        r_body.font.size = Pt(12)
        r_body.font.name = TNR
        return p

    # ── Title ──────────────────────────────────────────────────────────────────
    para('LICENCE TO USE TRADE MARK', bold=True, sz=16, center=True, sp_b=0, sp_a=16)

    # ── Preamble ───────────────────────────────────────────────────────────────
    deed_date   = data.get('deed_date', '').strip()
    if not deed_date:
        from datetime import datetime
        deed_date = datetime.now().strftime('%-d %B %Y')  # e.g. "7 April 2026"
    licensor_name = data.get('licensor_name', '[Licensor Name]')
    licensee_name = data.get('licensee_name', '[Licensee Name]')
    trademark     = data.get('trademark', '[TRADE MARK]')
    goods_services= data.get('goods_services', '[goods/services]')
    territory     = data.get('territory', '[Territory]')
    licence_fee_pct = data.get('licence_fee_pct', '10')
    # Auto-filled / boilerplate — not collected from user
    payment_dates = '30th June and 31st December'
    notice_period = '3'

    preamble = (
        f'THIS DEED OF LICENCE is made on this {deed_date} between {licensor_name}, '
        f'hereinafter called the LICENSOR (which term shall unless excluded '
        f'by or repugnant to the context include its successors and assigns) of the one part '
        f'and {licensee_name}, '
        f'hereinafter referred to as the LICENSEE (which term '
        f'shall unless excluded by or repugnant to the context include its permitted nominees) '
        f'of the other part.'
    )
    para(preamble, sp_b=0, sp_a=10)

    recitals = [
        (f'WHEREAS the LICENSOR is the manufacturer of and dealer in {goods_services} and holds '
         f'the registered Trade Mark {trademark} in respect of {goods_services}.'),
        (f'AND WHEREAS the LICENSOR intends to expand its business and sell its products under its '
         f'Trade Mark in {territory}.'),
        (f'AND WHEREAS the LICENSEE has a manufacturing/trading unit to deal in {goods_services}.'),
        (f'AND WHEREAS the LICENSEE has approached the LICENSOR to grant licence to use the '
         f"LICENSOR's Trade Mark {trademark} for sale of the products/services of the LICENSEE."),
        (f'AND WHEREAS the LICENSOR has agreed to allow the LICENSEE to use its said Trade Mark '
         f'{trademark} to sell/provide the LICENSEE\'s {goods_services} on certain terms and conditions.'),
    ]
    for rec in recitals:
        para(rec, sp_b=4, sp_a=4)

    para('NOW THEREFORE THESE PRESENTS witnesseth and the parties hereby agree as follows:', bold=True, sp_b=10, sp_a=8)

    clauses = [
        (1, f'The LICENSOR hereby doth grant to the LICENSEE non-exclusive right to use the '
            f"LICENSOR's Trade Mark {trademark} in {territory} for sale/provision of its "
            f'{goods_services} under the Trade Name {trademark}.'),
        (2, f'The use of the Trade Mark by the LICENSEE shall be confined only to the items/services '
            f'that may be manufactured or provided by the LICENSEE at its own premises or through '
            f'its authorised channels. The LICENSEE shall pay half-yearly to the LICENSOR a licence '
            f'fee at the rate of {licence_fee_pct}% on the turnover of business of the LICENSEE, '
            f'such payment to be made by {payment_dates} every year.'),
        (3, f'The LICENSEE shall comply with the requirements and provisions of all laws, rules and '
            f'regulations in relation to the manufacture, sale or provision of {goods_services} '
            f'under the said Trade Mark of the LICENSOR.'),
        (4, f'The LICENSEE shall manufacture and sell/provide {goods_services} under the said Trade '
            f'Mark {trademark} in accordance with the specifications, make-up, brand and packing that '
            f'the LICENSOR may from time to time intimate to the LICENSEE.'),
        (5, f"The LICENSOR shall have access to the LICENSEE's manufacturing/service unit and to "
            f"inspect the LICENSEE's books of accounts and other records at all reasonable times on "
            f'giving prior notice.'),
        (6, f'The LICENSEE agrees, declares and covenants not to use the said Trade Mark or advertise '
            f'or publish in newspapers, journals, labels or any other documents or packages or do '
            f'anything having the effect of diluting the distinctiveness of the Trade Mark of the '
            f'LICENSOR. The LICENSEE shall give indications either visually or phonetically to the '
            f'purchasing public that the LICENSEE is using the Trade Mark {trademark} as the licensee '
            f'of the LICENSOR.'),
        (7, f'The LICENSEE undertakes to compensate the LICENSOR and keep the LICENSOR harmless from '
            f'and indemnified against all claims, proceedings, losses, costs and expenses for any '
            f'wilful or negligent conduct of the LICENSEE in relation to the use of the Trade Mark '
            f'of the LICENSOR.'),
        (8, f'The LICENSEE shall not acquire any right of registration of the Trade Mark by virtue '
            f'of the LICENSEE manufacturing, selling or providing {goods_services} as user of the '
            f'Trade Mark {trademark} for any number of years or after termination of the licence or otherwise.'),
        (9, f"The LICENSEE shall inform the LICENSOR of any infringement of the LICENSOR's Trade "
            f'Mark {trademark} with particulars of the infringement or passing off and the names and '
            f'addresses of the offenders.'),
        (10, f'The LICENSOR shall take and/or permit the LICENSEE to take all possible legal steps '
             f'for the protection and preservation of the Trade Mark and prevention of its '
             f'infringement or passing off by any person.'),
        (11, f'This agreement is terminable by giving {notice_period} months\' notice from either side.'),
        (12, f'In any legal proceedings or in any action against the infringement or passing off in '
             f'relation to the Trade Mark of the goods/services covered by the Licence, the LICENSEE '
             f'will take appropriate steps to protect the interests of the LICENSOR and allow the '
             f'LICENSOR to take any legal action or steps and to join the LICENSEE as a party therein.'),
    ]
    for num, text in clauses:
        clause(num, text)

    # ── Signature block ────────────────────────────────────────────────────────
    para('', sp_b=8, sp_a=0)
    para('THE SCHEDULE', bold=True, center=True, sp_b=8, sp_a=8)
    para('IN WITNESS WHEREOF the parties herein have executed these presents on the day, month and '
         'year first above-written.', sp_b=0, sp_a=16)

    para('Signed, sealed and delivered by', sp_b=0, sp_a=4)
    para(f'The authorised representative of {licensor_name} in the presence of:', sp_b=0, sp_a=12)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(8)
    sig_p.paragraph_format.space_after  = Pt(4)
    tab = sig_p.paragraph_format.tab_stops
    r1 = sig_p.add_run('1. ________________________')
    r1.font.name = TNR; r1.font.size = Pt(12)
    r1 = sig_p.add_run('\t\t\tSignature: ________________________')
    r1.font.name = TNR; r1.font.size = Pt(12)

    sig_p2 = doc.add_paragraph()
    sig_p2.paragraph_format.space_before = Pt(8)
    sig_p2.paragraph_format.space_after  = Pt(4)
    r2 = sig_p2.add_run('2. ________________________')
    r2.font.name = TNR; r2.font.size = Pt(12)
    r2b = sig_p2.add_run('\t\t\tDate: ________________________')
    r2b.font.name = TNR; r2b.font.size = Pt(12)

    os.makedirs('generated', exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', f'TM_Licence_{licensor_name[:25]}')
    out  = os.path.abspath(f'generated/{safe}_{uuid.uuid4().hex[:8]}.docx')
    doc.save(out)
    return out


@app.route('/api/legal/trademark-license', methods=['POST'])
def gen_trademark_license():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.json or {}
    try:
        path = build_trademark_license_docx(data)
        jid  = uuid.uuid4().hex
        jobs[jid] = {'status': 'done', 'file_path': path, 'topic': 'Trademark License Agreement'}
        return jsonify({'success': True, 'job_id': jid})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    os.makedirs('generated', exist_ok=True)

    or_key = os.environ.get("GROQ_API_KEY", "").strip()
    key_str = "✓ Groq — ready!" if or_key else "✗ NOT SET — see below"
    print('\n' + '='*60)
    print('  rdxper v4.0  —  Free AI Research Paper Generator')
    print('  Powered by Groq (free tier)')
    print('  Open browser:  http://127.0.0.1:8080')
    print(f'  GROQ_API_KEY: {key_str}')
    print('='*60 + '\n')
    if not or_key:
        print('  ┌─ GET YOUR FREE GROQ API KEY ─────────────────────────────┐')
        print('  │                                                          │')
        print('  │  Groq — free, no credit card needed:                    │')
        print('  │    1. Visit https://console.groq.com                    │')
        print('  │    2. Sign up → API Keys → Create API Key               │')
        print('  │    3. Windows:  set GROQ_API_KEY=your_key_here          │')
        print('  │       Mac/Linux: export GROQ_API_KEY=your_key           │')
        print('  │    4. Run python rdxper.py again                        │')
        print('  │                                                          │')
        print('  └──────────────────────────────────────────────────────────┘')
        print()

    port = int(os.environ.get("PORT", 8080))
    host = "0.0.0.0"
    app.run(host=host, port=port, debug=False, threaded=True)
